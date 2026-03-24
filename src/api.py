import sys, io
if sys.stdout.encoding != "utf-8":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if sys.stderr.encoding != "utf-8":
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dotenv import load_dotenv
load_dotenv()  # Carga .env desde el directorio de trabajo

from fastapi import FastAPI, UploadFile, File, Header, HTTPException, Depends, Request, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.staticfiles import StaticFiles
from fastapi.exceptions import RequestValidationError
from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from pydantic import BaseModel, ValidationError, model_validator
from src.schemas import UserSchema, UserOut  # UserOut = alias de UserSchema
import bcrypt
from jose import jwt, JWTError
import os
import re
import unicodedata
import hashlib
import json
import traceback
from datetime import datetime, timezone, timedelta
import secrets
from pathlib import Path
from src.auth_utils import (
    create_access_token,
    enforce_internal_role,
    generate_hybrid_id,
    is_hybrid_id,
    JWT_SECRET as _AUTH_JWT_SECRET,
    JWT_ALGORITHM as _AUTH_JWT_ALGORITHM,
)
from google.cloud import storage
from google.auth.exceptions import DefaultCredentialsError
from google.api_core.exceptions import Forbidden, Unauthorized
from google.oauth2 import service_account
from src.adapters.google_cloud import GeminiAuditor
from src.adapters.document_ai import DocumentAIAdapter
from src.core.data_contract import build_contract
from src.core.db_writer import (
    insert_contract, ensure_schema, update_kpi_value,
    lookup_portfolio, detect_company_from_text, PORTFOLIO_MAP,
    query_portfolio_analytics, run_audit_query, audit_contract,
    run_fidelity_audit, COMPANY_BUCKET, insert_upload_log,
    insert_ai_audit_log, query_kpi_metadata, query_coverage,
)
from src.core.data_contract import build_checklist_status, KPI_REGISTRY
import pandas as pd

app = FastAPI(title="Cometa Pipeline API", version="1.0.0")

# ── A2: Rate limiting ──────────────────────────────────────────────────────────
limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# ── E1: Manejo global de errores de validación Pydantic → 422 ─────────────────
#
# Por qué dos handlers:
#   - ValidationError      → errores internos (UserSchema.model_validate, etc.)
#   - RequestValidationError → errores de FastAPI al parsear el body/query params
#
# Ambos devuelven la misma estructura {detail: [{loc, msg, type}]} para que el
# frontend pueda parsearlos de forma uniforme sin lógica de ramificación.
#
# GARANTÍA DE ORDEN: los handlers se registran aquí, antes de cualquier ruta.
# Python evalúa el decorador en el momento de la definición, así que estos
# handlers están activos desde el primer request — incluido cualquier intento
# de escritura que falle en UserSchema.model_validate().

def _format_validation_errors(errors: list[dict]) -> list[dict]:
    """
    Normaliza la lista de errores de Pydantic v2 al subset {loc, msg, type}.
    Excluye 'url', 'input' y 'ctx' que son ruido para el cliente.
    """
    return [
        {
            "loc":  list(e.get("loc", [])),
            "msg":  e.get("msg", ""),
            "type": e.get("type", ""),
        }
        for e in errors
    ]


@app.exception_handler(ValidationError)
async def pydantic_validation_handler(
    request: Request,
    exc: ValidationError,
) -> JSONResponse:
    """
    Captura pydantic.ValidationError lanzado dentro de cualquier route handler.
    Ejemplo: UserSchema.model_validate() falla → este handler retorna 422
    ANTES de que se abra ningún archivo para escritura.
    """
    return JSONResponse(
        status_code=422,
        content={"detail": _format_validation_errors(exc.errors(include_url=False))},
    )


@app.exception_handler(RequestValidationError)
async def request_validation_handler(
    request: Request,
    exc: RequestValidationError,
) -> JSONResponse:
    """
    Sobreescribe el handler por defecto de FastAPI para request body / query params.
    Misma estructura {detail} que pydantic_validation_handler → frontend unificado.
    """
    return JSONResponse(
        status_code=422,
        content={"detail": _format_validation_errors(exc.errors())},
    )


@app.on_event("startup")
async def _startup():
    """Bootstrap BigQuery tables once at server start."""
    import asyncio
    print("🚀 [Startup] Servidor iniciando — intentando bootstrap de BigQuery...")
    try:
        # ensure_schema() hace I/O bloqueante (BigQuery). Lo corremos en un
        # thread con timeout de 15 s para que el servidor nunca quede zombi
        # si GCP no responde (ej. error 403 o red no disponible).
        await asyncio.wait_for(
            asyncio.get_event_loop().run_in_executor(None, ensure_schema),
            timeout=15.0
        )
        print("✅ [Startup] BigQuery schema OK")
    except asyncio.TimeoutError:
        print("⚠️  [Startup] BigQuery schema bootstrap TIMEOUT (15 s) — servidor listo sin BQ")
    except Exception as e:
        # Non-fatal: if BQ is unreachable the upload still works via GCS.
        print(f"⚠️  [Startup] BigQuery schema bootstrap falló (non-fatal): {e}")
    print("✅ [Startup] Servidor listo para recibir archivos en :8000")

cors_origins_raw = os.getenv(
    "CORS_ORIGINS",
    json.dumps([
        # ── Production ────────────────────────────────────────────────────────
        "https://cometa-vault-frontend-92572839783.us-central1.run.app",
        # ── Local development ─────────────────────────────────────────────────
        "http://localhost:3000",
        "http://localhost:3001",
        "http://localhost:3002",
        "http://localhost:3003",
        "http://localhost:8000",
    ]),
)
try:
    cors_origins = json.loads(cors_origins_raw)
    if not isinstance(cors_origins, list):
        cors_origins = ["http://localhost:3000"]
except Exception:
    cors_origins = ["http://localhost:3000"]

app.add_middleware(
    CORSMiddleware,
    allow_origins=cors_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "OPTIONS"],
    allow_headers=["Authorization", "Content-Type", "founder-email", "company-id"],
)

# ═══════════════════════════════════════════════════════════════════════════════
# BLOQUE DE SEGURIDAD
# ═══════════════════════════════════════════════════════════════════════════════

# ── C3: JWT Authentication ────────────────────────────────────────────────────
_bearer_scheme = HTTPBearer(auto_error=False)
_JWT_SECRET    = _AUTH_JWT_SECRET   # from auth_utils → JWT_SECRET env var
_JWT_ALGORITHM = _AUTH_JWT_ALGORITHM

# ── Auth: ruta al fichero de usuarios ─────────────────────────────────────────
_USERS_FILE = Path(__file__).parent / "users.json"

async def _require_auth(
    credentials: HTTPAuthorizationCredentials = Depends(_bearer_scheme),
) -> dict:
    """
    Valida el JWT HS256 emitido por /api/auth/token (Next.js).
    Lanza 401 si el token es inválido, expirado o ausente.
    """
    if not credentials:
        raise HTTPException(status_code=401, detail="Token de autenticación requerido")
    if not _JWT_SECRET:
        raise HTTPException(status_code=500, detail="NEXTAUTH_SECRET no configurado en el servidor")
    try:
        payload = jwt.decode(
            credentials.credentials,
            _JWT_SECRET,
            algorithms=[_JWT_ALGORITHM],
            options={"verify_aud": False},
        )
        return payload
    except JWTError as exc:
        raise HTTPException(status_code=401, detail=f"Token inválido: {exc}")

# ── C7: Magic bytes validation ────────────────────────────────────────────────
_MAGIC_BYTES: dict[str, list[bytes]] = {
    ".pdf":     [b"%PDF"],
    ".xlsx":    [b"PK\x03\x04"],
    ".xls":     [b"\xd0\xcf\x11\xe0"],
    ".docx":    [b"PK\x03\x04"],
    ".doc":     [b"\xd0\xcf\x11\xe0"],
    ".parquet": [b"PAR1"],
    ".csv":     [],  # Texto plano — no tiene magic bytes fijos
}

def _validate_magic_bytes(file_content: bytes, ext: str) -> bool:
    """
    Verifica que los primeros bytes del contenido coincidan con la extensión declarada.
    Protege contra archivos renombrados (p.ej. malware.exe → informe.pdf).
    """
    signatures = _MAGIC_BYTES.get(ext, [])
    if not signatures:
        return True
    return any(file_content[:8].startswith(sig) for sig in signatures)

# ── C2: Límite de tamaño de archivo ──────────────────────────────────────────
_MAX_FILE_MB    = int(os.getenv("MAX_FILE_SIZE_MB", "50"))
_MAX_FILE_BYTES = _MAX_FILE_MB * 1024 * 1024

# ── C6: Sanitización de nombre de archivo ────────────────────────────────────
_SAFE_FILENAME_RE = re.compile(r"[^\w\-.]")

def _sanitize_filename(filename: str) -> str:
    """
    Protege contra path traversal y caracteres peligrosos en nombres de archivo.
    Pasos: normalizar unicode → extraer basename → eliminar chars no seguros
           → eliminar puntos dobles → limitar a 200 chars.
    """
    filename = unicodedata.normalize("NFKD", filename)
    filename = os.path.basename(filename)                     # Bloquea ../../
    filename = _SAFE_FILENAME_RE.sub("_", filename)           # Solo alfanum + -_.
    filename = re.sub(r"\.{2,}", ".", filename)               # Elimina ..
    stem, ext = os.path.splitext(filename)
    return f"{stem[:196]}{ext}" if len(filename) > 200 else filename

# ── C5: Validación de headers de entrada ─────────────────────────────────────
_COMPANY_ID_RE = re.compile(r"^[a-zA-Z0-9_\-\.]{1,64}$")

def _validate_company_header(company_id: str | None) -> str | None:
    """Valida que company_id sea alfanumérico + guiones/puntos (sin path traversal)."""
    if not company_id:
        return None
    if not _COMPANY_ID_RE.match(company_id):
        raise HTTPException(
            status_code=400,
            detail=f"company_id contiene caracteres no permitidos: {company_id!r}"
        )
    return company_id

def _validate_email_header(email: str | None) -> str | None:
    """Valida formato básico de email. No verifica entregabilidad."""
    if not email:
        return None
    _EMAIL_RE = re.compile(r"^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$")
    if not _EMAIL_RE.match(email):
        raise HTTPException(status_code=400, detail=f"founder-email inválido: {email!r}")
    return email.lower()

# ── C4: Verificación de origen / preparación Cloud IAP ───────────────────────
_SKIP_ORIGIN_CHECK = os.getenv("SKIP_ORIGIN_CHECK", "false").lower() == "true"
_INTERNAL_SOURCE_HEADER = "x-cometa-source"
_IAP_USER_HEADER = "x-goog-authenticated-user-email"
_VALID_COMETA_SOURCES = {"dashboard", "analyst-portal", "internal-tool"}

async def _verify_origin(request: Request) -> None:
    """
    C4: Verifica que la petición provenga de una fuente autorizada.
    En producción (Cloud IAP) valida el header X-Goog-Authenticated-User-Email.
    En entornos sin IAP, acepta el header X-Cometa-Source con valor válido.
    SKIP_ORIGIN_CHECK=true lo deshabilita para desarrollo local.
    """
    if _SKIP_ORIGIN_CHECK:
        return

    # Cloud IAP en producción inyecta este header automáticamente
    iap_user = request.headers.get(_IAP_USER_HEADER)
    if iap_user:
        return  # IAP verificó la identidad; request autorizada

    # Fallback para entornos sin IAP (staging interno, tests de integración)
    source = request.headers.get(_INTERNAL_SOURCE_HEADER, "").strip().lower()
    if source in _VALID_COMETA_SOURCES:
        return

    raise HTTPException(
        status_code=403,
        detail="Acceso denegado: origen no autorizado. Se requiere X-Goog-Authenticated-User-Email o X-Cometa-Source válido.",
    )

# ── A1: Derivación de tenant desde JWT ───────────────────────────────────────
_INTERNAL_DOMAINS = {"cometa.vc", "cometa.fund", "cometavc.com"}

def _derive_tenant_from_token(token: dict) -> str | None:
    """
    A1: Extrae company_id del dominio del email en el JWT.
    - Analistas internos (@cometa.vc, etc.) → None (pueden consultar cualquier empresa)
    - Founders externos → company_id derivado de su dominio de email
    El llamador NO puede sobreescribir esto con company_id del body.
    """
    email: str = token.get("email", "")
    if not email or "@" not in email:
        return None
    domain = email.split("@", 1)[1].lower()
    if domain in _INTERNAL_DOMAINS:
        return None  # Analista interno — sin restricción de tenant
    # Founder externo: derivar company_id canónico desde su dominio
    comp_id, _, _, _ = get_company_id(domain)
    return comp_id


# ═══════════════════════════════════════════════════════════════════════════════
# BLOQUE DE NORMALIZACIÓN DE CONTRATO DE DATOS
#   R1 — normalize_period()  →  fecha libre → PYYYYQxMyy
#   R2 — get_company_id()    →  nombre libre → COMP_XXX + fund_id + bucket_id
# ═══════════════════════════════════════════════════════════════════════════════

# ── R1: Tablas de traducción para normalización de períodos ───────────────────

_MONTH_TO_NUM: dict[str, str] = {
    # Inglés
    "january":"01","february":"02","march":"03","april":"04",
    "may":"05","june":"06","july":"07","august":"08",
    "september":"09","october":"10","november":"11","december":"12",
    # Español
    "enero":"01","febrero":"02","marzo":"03","abril":"04",
    "mayo":"05","junio":"06","julio":"07","agosto":"08",
    "septiembre":"09","octubre":"10","noviembre":"11","diciembre":"12",
    # Abreviaturas EN
    "jan":"01","feb":"02","mar":"03","apr":"04","jun":"06",
    "jul":"07","aug":"08","sep":"09","oct":"10","nov":"11","dec":"12",
}

_MONTH_TO_QUARTER: dict[str, str] = {
    "01":"Q1","02":"Q1","03":"Q1",
    "04":"Q2","05":"Q2","06":"Q2",
    "07":"Q3","08":"Q3","09":"Q3",
    "10":"Q4","11":"Q4","12":"Q4",
}

# Primer mes canónico de cada quarter (para cuando solo tenemos Q sin mes exacto)
_QUARTER_FIRST_MONTH: dict[str, str] = {
    "1":"01","2":"04","3":"07","4":"10",
}

# Período de cierre de cada semestre (H1 → jun, H2 → dic)
_HALF_CLOSE_MONTH: dict[str, str] = {"1":"06","2":"12"}
_HALF_QUARTER: dict[str, str]     = {"1":"Q2","2":"Q4"}

# Período ya en formato canónico — pasa sin transformación
_CANONICAL_RE = re.compile(r"^P(20\d{2})Q([1-4])M(\d{2})$")


def normalize_period(date_str: str) -> tuple[str, bool]:
    """
    R1 — Normaliza cualquier representación de período al formato PYYYYQxMyy.

    Entradas reconocidas (insensibles a mayúsculas/espacios extra):
      "March 2025"    → "P2025Q1M03"
      "Q1 2025"       → "P2025Q1M01"
      "Q4 2024"       → "P2024Q4M10"
      "H1 2025"       → "P2025Q2M06"
      "H2 2024"       → "P2024Q4M12"
      "FY2025"        → "P2025Q4M12"
      "2025"          → "P2025Q4M12"
      "2025M03"       → "P2025Q1M03"
      "P2025Q1M03"    → "P2025Q1M03"  (passthrough)

    Returns
    -------
    (canonical_period_id, is_valid)
      is_valid=False  →  el input no pudo mapearse; se devuelve un fallback
                        con el año actual para no romper el pipeline.
    """
    if not date_str or not isinstance(date_str, str):
        fallback = f"P{datetime.now(timezone.utc).year}Q4M12"
        return fallback, False

    s = date_str.strip()

    # 0. Ya canónico — passthrough sin coste
    if _CANONICAL_RE.match(s):
        return s, True

    sl = s.lower()

    # 1. "March 2025" / "marzo 2025" — nombre de mes + año
    for month_name, month_num in _MONTH_TO_NUM.items():
        m = re.search(rf"\b{re.escape(month_name)}\b\s*(20\d{{2}})", sl)
        if not m:
            # También acepta "2025 March"
            m = re.search(rf"(20\d{{2}})\s*\b{re.escape(month_name)}\b", sl)
        if m:
            year = m.group(1) if m.lastindex == 1 else (m.group(1) if m.group(1).startswith("20") else m.group(2))
            # Reparar: si el grupo de año no es el que tiene 20xx, tomar el otro
            year = next((g for g in m.groups() if g and g.startswith("20")), None)
            if not year:
                continue
            quarter = _MONTH_TO_QUARTER[month_num]
            return f"P{year}{quarter}M{month_num}", True

    # 2. "Q1 2025" / "Q4 2024"
    m = re.search(r"q([1-4])\s*[/\-]?\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"(20\d{2})\s*[/\-]?\s*q([1-4])", sl)
    if m:
        groups = m.groups()
        if groups[0].startswith("20"):
            year, qnum = groups[0], groups[1]
        else:
            qnum, year = groups[0], groups[1]
        return f"P{year}Q{qnum}M{_QUARTER_FIRST_MONTH[qnum]}", True

    # 3. "H1 2025" / "H2 2024"
    m = re.search(r"h([12])\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"(20\d{2})\s*h([12])", sl)
    if m:
        groups = m.groups()
        if groups[0].startswith("20"):
            year, half = groups[0], groups[1]
        else:
            half, year = groups[0], groups[1]
        return f"P{year}{_HALF_QUARTER[half]}M{_HALF_CLOSE_MONTH[half]}", True

    # 4. "FY2025" / "FY 2025" / "fiscal year 2025"
    m = re.search(r"fy\s*(20\d{2})", sl)
    if not m:
        m = re.search(r"fiscal\s+year\s*(20\d{2})", sl)
    if m:
        year = m.group(1)
        return f"P{year}Q4M12", True

    # 5. "2025M03" (formato partial canónico sin prefijo P)
    m = re.match(r"^(20\d{2})m(\d{2})$", sl)
    if m:
        year, month_num = m.group(1), m.group(2).zfill(2)
        quarter = _MONTH_TO_QUARTER.get(month_num, "Q4")
        return f"P{year}{quarter}M{month_num}", True

    # 6. Año suelto: "2025"
    m = re.fullmatch(r"20\d{2}", s.strip())
    if m:
        return f"P{s.strip()}Q4M12", True

    # 7. Cualquier año 20xx detectado en el string — fallback degradado
    m = re.search(r"(20\d{2})", s)
    if m:
        year = m.group(1)
        return f"P{year}Q4M12", False

    # 8. Sin información → año corriente, marcar como inválido
    fallback = f"P{datetime.now(timezone.utc).year}Q4M12"
    return fallback, False


# ── R2: Catálogo maestro company_key → COMP_XXX ──────────────────────────────
# Generado automáticamente desde PORTFOLIO_MAP + COMPANY_BUCKET.
# Se construye una vez al inicio del módulo para evitar recalcularlo por request.

def _build_company_catalog() -> dict[str, dict]:
    """
    Construye el catálogo: company_key → {comp_id, fund_id, bucket_id}.
    Importación tardía de PORTFOLIO_MAP/COMPANY_BUCKET para evitar
    que este módulo cargue antes de que db_writer esté inicializado.
    """
    from src.core.db_writer import PORTFOLIO_MAP, COMPANY_BUCKET
    catalog: dict[str, dict] = {}
    for key, info in PORTFOLIO_MAP.items():
        bucket = COMPANY_BUCKET.get(key, "OTH")
        comp_id = f"COMP_{key.upper().replace('-','_').replace('.','_')}"
        catalog[key] = {
            "comp_id":  comp_id,
            "fund_id":  info["portfolio_id"],
            "bucket_id": bucket,
        }
    return catalog

_COMPANY_CATALOG: dict[str, dict] = {}   # lazy init en primer uso


def _get_company_catalog() -> dict[str, dict]:
    global _COMPANY_CATALOG
    if not _COMPANY_CATALOG:
        _COMPANY_CATALOG = _build_company_catalog()
    return _COMPANY_CATALOG


def get_company_id(name_str: str) -> tuple[str, str, str, bool]:
    """
    R2 — Mapea texto libre al ID canónico COMP_XXX y hereda fund_id / bucket_id
    desde el catálogo maestro (PORTFOLIO_MAP + COMPANY_BUCKET).

    Estrategia de resolución (en orden de precisión, para en el primer match):
      1. Normalizar: minúsculas, strip dominio (.com/.mx/etc.)
      2. Match exacto contra catalog keys
      3. Strip guiones/guiones bajos y match exacto
      4. Prefijo: "m1-insurtech" → "m1"
      5. Substring: catalog_key contenido en el input normalizado
      6. Input contenido como substring de alguna catalog_key

    Returns
    -------
    (canonical_id, fund_id, bucket_id, is_known)
      is_known=False →  empresa no registrada; comp_id es COMP_UNKNOWN_<hash>
                        para garantizar trazabilidad sin perder la submission.
    """
    import hashlib as _hashlib

    if not name_str:
        return "COMP_UNKNOWN", "unknown", "OTH", False

    catalog = _get_company_catalog()

    # ── Normalización base ────────────────────────────────────────────────
    base = name_str.lower().strip()
    base = base.split(".")[0]                        # strip TLD: "simetrik.com" → "simetrik"
    base = re.sub(r"\s+", " ", base)                # colapsar espacios

    def _entry(key: str) -> tuple[str, str, str, bool]:
        e = catalog[key]
        return e["comp_id"], e["fund_id"], e["bucket_id"], True

    # 1. Exacto
    if base in catalog:
        return _entry(base)

    # 2. Strip separadores
    stripped = re.sub(r"[-_\s]", "", base)
    for key in catalog:
        if re.sub(r"[-_\s]", "", key) == stripped:
            return _entry(key)

    # 3. Prefijo: "m1-insurtech" starts with "m1-" o "m1_"
    for key in sorted(catalog.keys(), key=len, reverse=True):
        if base.startswith(key + "-") or base.startswith(key + "_") or base.startswith(key + " "):
            return _entry(key)

    # 4. El input contiene el key como palabra/token
    for key in sorted(catalog.keys(), key=len, reverse=True):
        pattern = r"(?<![a-z])" + re.escape(key) + r"(?![a-z])"
        if re.search(pattern, base):
            return _entry(key)

    # 5. El key contiene el input (abreviaciones: "solvento" en "solvento financiero")
    for key in sorted(catalog.keys(), key=len, reverse=True):
        if stripped and stripped in re.sub(r"[-_\s]", "", key):
            return _entry(key)

    # Sin match → COMP_UNKNOWN_<fingerprint> (determinístico, trazable)
    fingerprint = _hashlib.sha1(name_str.lower().encode()).hexdigest()[:8].upper()
    unknown_id  = f"COMP_UNKNOWN_{fingerprint}"
    print(f"⚠️  [R2] '{name_str}' no está en el catálogo maestro → {unknown_id}")
    return unknown_id, "unknown", "OTH", False


def _apply_contract_normalization(
    contract: dict,
    raw_company: str,
    raw_period: str,
) -> dict:
    """
    Aplica R1 y R2 al contrato ya construido por build_contract().

    Muta el contrato in-place y devuelve un dict con los resultados
    de normalización para logging y trazabilidad.

    Regla de oro:
      Si period o company no se pueden mapear → submission.status = "error"
      El contrato SE MUTA en todo caso para garantizar que el pipeline
      continúa y no pierde datos: is_known/period_valid se registran en
      submission para que el analista pueda corregir manualmente.
    """
    # ── R1: Normalizar período ────────────────────────────────────────────
    # Prioridad: raw_period del _document_context de Gemini (más específico)
    # → contiene "March 2025", "Q1 2025", etc. que infer_period_id() ignora.
    # Fallback: el period_id ya inferido por build_contract() (solo tiene año).
    current_period = raw_period or contract["submission"].get("period_id", "")
    norm_period, period_ok = normalize_period(current_period)

    # ── R2: Normalizar company_id ─────────────────────────────────────────
    comp_id, fund_id, bucket_id, company_ok = get_company_id(raw_company)

    # ── Determinar status final ───────────────────────────────────────────
    normalization_errors: list[str] = []
    if not period_ok:
        normalization_errors.append(
            f"period '{current_period}' no reconocido — "
            f"se usó fallback '{norm_period}'"
        )
    if not company_ok:
        normalization_errors.append(
            f"company '{raw_company}' no está en el catálogo maestro — "
            f"se asignó '{comp_id}'"
        )

    # Solo marca error si ambos fallan simultáneamente (un error de compañía
    # con período conocido es recuperable vía edición manual del analista)
    if not period_ok and not company_ok:
        contract["submission"]["status"] = "error"
    elif not company_ok:
        contract["submission"]["status"] = "pending_review"

    # ── Mutar submission ──────────────────────────────────────────────────
    contract["submission"]["period_id"]            = norm_period
    contract["submission"]["company_id"]           = comp_id
    contract["submission"]["fund_id"]              = fund_id
    contract["submission"]["bucket_id"]            = bucket_id
    contract["submission"]["period_normalized"]    = period_ok
    contract["submission"]["company_known"]        = company_ok
    if normalization_errors:
        contract["submission"]["normalization_errors"] = normalization_errors

    # ── Mutar kpi_rows — propagar IDs canónicos a cada hecho ─────────────
    for row in contract["kpi_rows"]:
        row["period_id"]  = norm_period   # sobrescribe el inferido por build_contract
        row["company_id"] = comp_id       # añade company_id a la fila (requerido por contrato)
        row["fund_id"]    = fund_id

    return {
        "period_id":    norm_period,
        "period_ok":    period_ok,
        "comp_id":      comp_id,
        "fund_id":      fund_id,
        "bucket_id":    bucket_id,
        "company_ok":   company_ok,
        "errors":       normalization_errors,
    }


# ── Rutas de archivos estáticos ───────────────────────────────────────────────
# BASE_DIR apunta a la raíz del proyecto (/app en Cloud Run) sin importar
# desde qué directorio se invoque uvicorn. Evita problemas con rutas relativas.
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_assets_dir = os.path.join(BASE_DIR, "assets")
if os.path.isdir(_assets_dir):
    app.mount("/assets", StaticFiles(directory=_assets_dir), name="assets")

@app.get("/api/v2/health", include_in_schema=False)
async def root():
    return {"status": "backend_online", "version": "2.0"}

# Configuración
PROJECT_ID = os.getenv("GOOGLE_PROJECT_ID", "cometa-mvp")
LOCATION_DOC_AI = os.getenv("DOCUMENT_AI_LOCATION", "us")
PROCESSOR_ID = os.getenv("DOCUMENT_AI_PROCESSOR_ID", "c5e1adfde68e63cf")
VERTEX_LOCATION = os.getenv("VERTEX_AI_LOCATION", "us-central1")
GCS_INPUT_BUCKET = os.getenv("GCS_INPUT_BUCKET", "ingesta-financiera-raw-cometa-mvp")
GCS_OUTPUT_BUCKET = os.getenv("GCS_OUTPUT_BUCKET", "ingesta-financiera-raw-cometa-mvp")

def _resolve_service_account_path() -> str | None:
    env_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
    if env_path:
        return env_path
    # Fallback al JSON en raíz del repo
    repo_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    fallback = os.path.join(repo_root, "cometa_key.json")
    return fallback if os.path.exists(fallback) else None

def _parse_sa_json(raw: str, tag: str = "") -> dict:
    """
    Parsea el JSON de una service account desde una variable de entorno,
    tolerando: espacios/newlines extra y doble-serialización (string dentro de string).
    Lanza ValueError con diagnóstico si el resultado no tiene los campos requeridos.
    """
    raw = raw.strip()
    parsed = json.loads(raw)
    # Doble-serialización: el valor del secreto era una cadena JSON (string de string)
    if isinstance(parsed, str):
        print(f"⚠️  {tag} GCP_SERVICE_ACCOUNT_JSON estaba doblemente serializado — decodificando de nuevo")
        parsed = json.loads(parsed)
    if not isinstance(parsed, dict):
        raise ValueError(f"{tag} GCP_SERVICE_ACCOUNT_JSON no es un objeto JSON válido (tipo: {type(parsed).__name__})")
    required = {"type", "project_id", "private_key", "client_email"}
    missing = required - parsed.keys()
    if missing:
        raise ValueError(
            f"{tag} GCP_SERVICE_ACCOUNT_JSON le faltan campos: {missing}. "
            f"Claves presentes: {list(parsed.keys())}"
        )
    print(f"✅  {tag} Service account JSON OK — client_email: {parsed.get('client_email')}")
    return parsed


def _load_gcp_credentials():
    # ── Prioridad 1: GCP_SERVICE_ACCOUNT_JSON (Cloud Run + Secret Manager) ──────
    # El JSON completo se inyecta como variable de entorno desde Secret Manager.
    # No requiere archivo físico — compatible con contenedores inmutables.
    sa_json_str = os.getenv("GCP_SERVICE_ACCOUNT_JSON")
    if sa_json_str:
        print("🔐 [GCP] Usando GCP_SERVICE_ACCOUNT_JSON (Secret Manager)")
        sa_info = _parse_sa_json(sa_json_str, "[GCP]")
        creds = service_account.Credentials.from_service_account_info(sa_info)
        creds_project = getattr(creds, "project_id", None)
        if creds_project and creds_project != PROJECT_ID:
            print(f"⚠️  [GCP] project_id del JSON ({creds_project}) != PROJECT_ID ({PROJECT_ID})")
        return creds

    # ── Prioridad 2: GOOGLE_APPLICATION_CREDENTIALS (archivo, dev local) ────────
    sa_path = _resolve_service_account_path()
    if not sa_path:
        raise DefaultCredentialsError(
            "No se encontró GCP_SERVICE_ACCOUNT_JSON ni GOOGLE_APPLICATION_CREDENTIALS "
            "y no existe cometa_key.json en la raíz"
        )
    if not os.path.isabs(sa_path):
        sa_path = os.path.abspath(sa_path)

    print(f"🔐 [GCP] Usando Service Account JSON: {sa_path}")
    if not os.path.exists(sa_path):
        raise DefaultCredentialsError(f"Service Account JSON no existe en: {sa_path}")

    creds = service_account.Credentials.from_service_account_file(sa_path)
    creds_project = getattr(creds, "project_id", None)
    if creds_project and creds_project != PROJECT_ID:
        print(
            f"⚠️  [GCP] project_id del JSON ({creds_project}) no coincide con PROJECT_ID ({PROJECT_ID})"
        )
    return creds

def _get_storage_client() -> storage.Client:
    """Crea un Storage client usando credenciales explícitas cuando es posible."""
    try:
        creds = _load_gcp_credentials()
        return storage.Client(project=PROJECT_ID, credentials=creds)
    except Exception as e:
        print(f"❌ [GCP] No se pudieron cargar credenciales explícitas: {e}")
        # Intentar fallback a ADC (podría funcionar en Cloud Run/GCE)
        return storage.Client(project=PROJECT_ID)

def get_file_hash(file_content: bytes) -> str:
    """Genera hash SHA-256 del contenido del archivo"""
    return hashlib.sha256(file_content).hexdigest()[:16]


def _is_financial_document(resultado: dict) -> bool:
    """
    Devuelve True si Gemini extrajo al menos 1 KPI financiero con valor real.
    Cualquiera de estos campos con un valor no-nulo califica el documento.
    Usado como gate antes de persistir en GCS / BigQuery.
    """
    fm = resultado.get("financial_metrics_2025")
    if not fm or not isinstance(fm, dict):
        return False

    # Rutas a los KPIs "core" — basta con que uno sea no-nulo
    _SENTINEL = {"", "null", "n/a", "--", "0", "none"}
    core_paths = [
        ["revenue_growth", "value"],
        ["base_metrics", "revenue", "value"],
        ["base_metrics", "ebitda", "value"],
        ["profit_margins", "gross_profit_margin", "value"],
        ["profit_margins", "ebitda_margin", "value"],
        ["cash_flow_indicators", "cash_in_bank_end_of_year", "value"],
        ["cash_flow_indicators", "annual_cash_flow", "value"],
        ["debt_ratios", "working_capital_debt", "value"],
        ["sector_metrics", "mrr", "value"],
        ["sector_metrics", "gmv", "value"],
        ["sector_metrics", "portfolio_size", "value"],
        ["sector_metrics", "loss_ratio", "value"],
    ]
    for path in core_paths:
        node = fm
        for key in path:
            if not isinstance(node, dict):
                node = None
                break
            node = node.get(key)
        if node is not None and str(node).strip().lower() not in _SENTINEL:
            return True
    return False


def _extract_kpi_confidence_scores(resultado: dict) -> dict[str, int]:
    """
    Extract per-KPI confidence scores from a parsed Gemini JSON result.

    Traverses every KPI path defined in KPI_REGISTRY. If a node has a
    ``confidence`` field (float 0.0–1.0 as instructed in the prompt), it is
    converted to an integer 0–100 and stored under the ``kpi_key``.

    KPIs without a value (null) or without a ``confidence`` field are omitted.

    Parameters
    ----------
    resultado : dict
        Parsed Gemini JSON (the ``resultado`` dict produced by the upload pipeline).

    Returns
    -------
    dict[str, int]
        Mapping of ``kpi_key`` → confidence integer (0–100).
        Empty dict if no confidence scores are available.
    """
    scores: dict[str, int] = {}
    for kpi_def in KPI_REGISTRY:
        path: list[str] = kpi_def["path"]
        kpi_key: str    = kpi_def["kpi_key"]

        node = resultado
        for segment in path:
            if not isinstance(node, dict):
                node = None
                break
            node = node.get(segment)

        if not isinstance(node, dict):
            continue

        raw_conf = node.get("confidence")
        if raw_conf is None:
            continue

        try:
            # Gemini returns float 0.0–1.0; convert to 0–100 integer
            conf_float = float(raw_conf)
            # Accept both 0-1 and 0-100 ranges defensively
            if conf_float <= 1.0:
                conf_int = round(conf_float * 100)
            else:
                conf_int = round(conf_float)
            scores[kpi_key] = max(0, min(100, conf_int))
        except (TypeError, ValueError):
            continue

    return scores


# ── RAG / Chat helpers ────────────────────────────────────────────────────────

def _get_bq_client_for_api():
    """BigQuery client using the same credential chain as the rest of the API."""
    from google.cloud import bigquery as bq
    try:
        creds = _load_gcp_credentials()
        return bq.Client(project=PROJECT_ID, credentials=creds)
    except Exception:
        return bq.Client(project=PROJECT_ID)


def _build_results_from_bq(company_id: str) -> list[dict]:
    """
    BigQuery fallback for /api/results.

    Reads fact_kpi_values for a company and synthesises AnalysisResult objects
    (one per distinct period_id) with a financial_metrics_2025 payload — the
    same structure the frontend's extractKPIs() already knows how to consume.

    Parameters
    ----------
    company_id : Lowercase canonical slug, e.g. ``"quinio"``.

    Returns
    -------
    list[dict]  — Empty list on any BQ error (never raises).
    """
    # kpi_key → (financial_metrics_2025 section, sub-key)
    # All known aliases from fact_kpi_values are listed; cash_at_hand is an
    # alternate column name used in some legacy loads alongside cash_in_bank_end_of_year.
    _KPI_PATH: dict[str, tuple[str, str]] = {
        "revenue":                  ("revenue",              "total_revenue"),
        "ebitda":                   ("income",               "net_income"),
        "net_income":               ("income",               "net_income"),
        "gross_profit_margin":      ("profit_margins",       "gross_profit_margin"),
        "gross_margin":             ("profit_margins",       "gross_profit_margin"),  # alias
        "ebitda_margin":            ("profit_margins",       "ebitda_margin"),
        "revenue_growth":           ("revenue_growth",       "value"),
        "cash_in_bank_end_of_year": ("cash_flow_indicators", "cash_in_bank_end_of_year"),
        "cash_at_hand":             ("cash_flow_indicators", "cash_in_bank_end_of_year"),  # alias
        "annual_cash_flow":         ("cash_flow_indicators", "annual_cash_flow"),
        "cogs":                     ("cost_structure",       "cogs"),
        "working_capital_debt":     ("debt_ratios",          "working_capital_debt"),
        "net_working_capital":      ("debt_ratios",          "net_working_capital"),
        "mrr":                      ("revenue",              "mrr"),
        "gmv":                      ("revenue",              "gmv"),
    }

    def _fmt(value: float, unit: str) -> str:
        """Format a raw numeric value into a display string matching the frontend's toK() expectations."""
        if unit == "%":
            return f"{value:.1f}%"
        if abs(value) >= 1_000_000:
            return f"${value / 1_000_000:.1f}M"
        if abs(value) >= 1_000:
            return f"${value / 1_000:.0f}K"
        return f"{value:,.2f}"

    try:
        from google.cloud import bigquery as _bq
        _client = _get_bq_client_for_api()
        _ds     = os.getenv("BIGQUERY_DATASET", "cometa_vault_test")

        sql = f"""
            SELECT
                submission_id,
                period_id,
                kpi_key,
                COALESCE(normalized_value_usd, numeric_value) AS num_value,
                raw_value,
                unit,
                value_status
            FROM `{PROJECT_ID}.cometa_vault.stg_legacy_fact_kpis`
            WHERE LOWER(company_id) = LOWER(@company_id)
              AND value_status IN ('legacy', 'verified')
            ORDER BY period_id ASC, kpi_key
        """
        job  = _client.query(
            sql,
            job_config=_bq.QueryJobConfig(query_parameters=[
                _bq.ScalarQueryParameter("company_id", "STRING", company_id)
            ])
        )
        rows = list(job.result())

        if not rows:
            return []

        # Group rows by period_id
        from collections import defaultdict
        by_period: dict[str, list] = defaultdict(list)
        for r in rows:
            by_period[r.period_id].append(r)

        # ── First pass: raw numeric values per period (for derived calculations) ──
        raw_by_period: dict[str, dict[str, float]] = {}
        for pid, period_rows in by_period.items():
            raw_by_period[pid] = {}
            for r in period_rows:
                if r.num_value is not None and r.kpi_key:
                    raw_by_period[pid][r.kpi_key] = float(r.num_value)

        sorted_periods = sorted(by_period.keys())

        results = []
        for i, period_id in enumerate(sorted_periods):
            period_rows = by_period[period_id]
            raw_vals    = raw_by_period[period_id]

            # Build financial_metrics_2025 from flat BQ rows
            fm: dict = {}
            for r in period_rows:
                if r.num_value is None:
                    continue  # skip nulls — leave section absent so frontend shows "—"
                path = _KPI_PATH.get(r.kpi_key)
                if not path:
                    continue
                section, subkey = path
                fm.setdefault(section, {})
                # Don't overwrite a key already populated by an earlier alias
                if subkey != "value" and subkey in fm[section]:
                    continue
                unit_str    = r.unit or ""
                display_val = _fmt(float(r.num_value), unit_str)
                if subkey == "value":
                    fm[section]["value"] = {"value": display_val, "unit": unit_str}
                else:
                    fm[section][subkey] = {"value": display_val, "unit": unit_str}

            # ── Derived: Revenue Growth from consecutive periods ──────────────
            if "value" not in fm.get("revenue_growth", {}):
                if i > 0:
                    rev_now  = raw_vals.get("revenue")
                    rev_prev = raw_by_period[sorted_periods[i - 1]].get("revenue")
                    if rev_now is not None and rev_prev and rev_prev != 0:
                        growth = (rev_now - rev_prev) / abs(rev_prev) * 100
                        fm.setdefault("revenue_growth", {})["value"] = {
                            "value": f"{growth:.1f}%", "unit": "%",
                        }

            # ── Derived: Gross Margin = (revenue − cogs) / revenue ────────────
            if "gross_profit_margin" not in fm.get("profit_margins", {}):
                rev  = raw_vals.get("revenue")
                cogs = raw_vals.get("cogs")
                if rev and rev != 0 and cogs is not None:
                    gm = (rev - cogs) / rev * 100
                    fm.setdefault("profit_margins", {})["gross_profit_margin"] = {
                        "value": f"{gm:.1f}%", "unit": "%",
                    }

            # ── Derived: EBITDA Margin = ebitda / revenue ─────────────────────
            if "ebitda_margin" not in fm.get("profit_margins", {}):
                rev    = raw_vals.get("revenue")
                ebitda = raw_vals.get("ebitda")
                if rev and rev != 0 and ebitda is not None:
                    em = ebitda / rev * 100
                    fm.setdefault("profit_margins", {})["ebitda_margin"] = {
                        "value": f"{em:.1f}%", "unit": "%",
                    }

            # ── Fidelity calculation ───────────────────────────────────────
            # A row is "verified" when the view returns value_status='verified'
            # (driven by is_manually_edited=TRUE in fact_kpi_values).
            non_null_rows   = [r for r in period_rows if r.num_value is not None]
            verified_count  = sum(1 for r in non_null_rows if r.value_status == "verified")
            total_count     = len(non_null_rows)
            fidelity_pct    = int(verified_count / total_count * 100) if total_count else 0
            period_status   = "verified" if fidelity_pct == 100 else "legacy"
            submission_ids  = list({r.submission_id for r in period_rows if r.submission_id})

            result_item = {
                "id":   f"legacy_{period_id}",
                "date": period_id,
                "data": {
                    "financial_metrics_2025": fm,
                    "_source":          "bigquery_legacy",
                    "_period_id":       period_id,
                    "_value_status":    period_status,
                    "_fidelity_pct":    fidelity_pct,
                    "_submission_ids":  submission_ids,
                },
                "metadata": {
                    "original_filename": f"histórico {period_id}",
                    "founder_email":     "",
                    "file_hash":         "",
                    "processed_at":      period_id,
                    "gcs_path":          "",
                    "company_domain":    company_id,
                    "portfolio_id":      "",
                },
            }
            results.append(result_item)

        print(f" [API/BQ] Resultados históricos para '{company_id}': "
              f"{len(results)} periodos, {len(rows)} filas")
        return results

    except Exception as _err:
        print(f" [API/BQ] Fallback BQ failed for '{company_id}' (non-fatal): {_err}")
        return []


def _build_all_results_from_bq(gcs_companies: set[str]) -> list[dict]:
    """
    Returns ONE synthetic AnalysisResult per company (latest available period)
    for companies NOT already covered by gcs_companies.

    KPI values are formatted for the extractTopKpis() path in the portfolio page:
      revenue_growth    → revenue_growth.value.value  (formatted as "X.X%")
      gross_profit_margin → profit_margins.gross_profit_margin.value
      ebitda_margin     → profit_margins.ebitda_margin.value

    Revenue Growth is calculated as:
      If revenue_growth kpi_key present → use directly.
      Otherwise → (revenue_last - revenue_prev) / |revenue_prev| * 100.

    Returns empty list on any error (never raises).
    """
    _KPI_PATH: dict[str, tuple[str, str]] = {
        "revenue":                  ("revenue",              "total_revenue"),
        "ebitda":                   ("income",               "net_income"),
        "net_income":               ("income",               "net_income"),
        "gross_profit_margin":      ("profit_margins",       "gross_profit_margin"),
        "gross_margin":             ("profit_margins",       "gross_profit_margin"),  # alias
        "ebitda_margin":            ("profit_margins",       "ebitda_margin"),
        "revenue_growth":           ("revenue_growth",       "value"),
        "cash_in_bank_end_of_year": ("cash_flow_indicators", "cash_in_bank_end_of_year"),
        "cash_at_hand":             ("cash_flow_indicators", "cash_in_bank_end_of_year"),  # alias
        "annual_cash_flow":         ("cash_flow_indicators", "annual_cash_flow"),
        "cogs":                     ("cost_structure",       "cogs"),
        "working_capital_debt":     ("debt_ratios",          "working_capital_debt"),
        "net_working_capital":      ("debt_ratios",          "net_working_capital"),
        "mrr":                      ("revenue",              "mrr"),
        "gmv":                      ("revenue",              "gmv"),
    }

    def _fmt(value: float, unit: str) -> str:
        if unit == "%":
            return f"{value:.1f}%"
        if abs(value) >= 1_000_000:
            return f"${value / 1_000_000:.1f}M"
        if abs(value) >= 1_000:
            return f"${value / 1_000:.0f}K"
        return f"{value:,.2f}"

    try:
        from collections import defaultdict as _dd
        _client = _get_bq_client_for_api()
        _ds     = os.getenv("BIGQUERY_DATASET", "cometa_vault_test")

        # Latest period per company + one-prior period for revenue growth calc
        sql = f"""
            WITH ranked AS (
                SELECT
                    company_id,
                    period_id,
                    kpi_key,
                    COALESCE(normalized_value_usd, numeric_value) AS num_value,
                    unit,
                    value_status,
                    ROW_NUMBER() OVER (
                        PARTITION BY company_id, kpi_key
                        ORDER BY period_id DESC
                    ) AS rn
                FROM `{PROJECT_ID}.cometa_vault.stg_legacy_fact_kpis`
                WHERE value_status IN ('legacy', 'missing_legacy', 'verified')
            )
            SELECT company_id, period_id, kpi_key, num_value, unit, value_status, rn
            FROM ranked
            WHERE rn <= 2
            ORDER BY company_id, kpi_key, rn
        """
        rows = list(_client.query(sql).result())

        # Separate: rn=1 (latest) and rn=2 (prior) per company+kpi
        latest: dict[tuple, float | None] = {}   # (company, kpi) → value
        prior:  dict[tuple, float | None] = {}
        period_map: dict[str, str]         = {}   # company → latest period_id

        for r in rows:
            # missing_legacy rows mark expected-but-absent values — exclude from calcs
            if r.value_status == "missing_legacy" or r.num_value is None:
                continue
            key = (r.company_id, r.kpi_key)
            if r.rn == 1:
                latest[key]              = r.num_value
                period_map[r.company_id] = r.period_id
            else:
                prior[key] = r.num_value

        # Build one result per company
        all_companies = {r.company_id for r in rows}
        results: list[dict] = []

        for company_id in sorted(all_companies):
            if company_id.lower() in gcs_companies:
                continue   # already covered by GCS

            period_id = period_map.get(company_id, "")

            # ── financial_metrics_2025 payload ───────────────────────────────
            fm: dict = {}
            for kpi_key, (section, subkey) in _KPI_PATH.items():
                val = latest.get((company_id, kpi_key))
                if val is None:
                    continue
                unit_key = next(
                    (r.unit for r in rows if r.company_id == company_id and r.kpi_key == kpi_key and r.rn == 1),
                    ""
                ) or ""

                # Revenue Growth: prefer stored value, else compute from revenue
                if kpi_key == "revenue_growth":
                    display = _fmt(val, "%")
                else:
                    display = _fmt(val, unit_key)

                fm.setdefault(section, {})
                if subkey == "value":
                    fm[section]["value"] = {"value": display, "unit": unit_key}
                else:
                    fm[section][subkey] = {"value": display, "unit": unit_key}

            # Derived Revenue Growth if kpi not stored
            if "revenue_growth" not in fm.get("revenue_growth", {}):
                rev_now  = latest.get((company_id, "revenue"))
                rev_prev = prior.get((company_id,  "revenue"))
                if rev_now is not None and rev_prev and rev_prev != 0:
                    growth = (rev_now - rev_prev) / abs(rev_prev) * 100
                    fm.setdefault("revenue_growth", {})["value"] = {
                        "value": f"{growth:.1f}%", "unit": "%"
                    }

            # Derived Gross Margin from revenue + cogs if not stored
            if "gross_profit_margin" not in fm.get("profit_margins", {}):
                rev  = latest.get((company_id, "revenue"))
                cogs = latest.get((company_id, "cogs"))
                if rev and rev != 0 and cogs is not None:
                    gm = (rev - cogs) / rev * 100
                    fm.setdefault("profit_margins", {})["gross_profit_margin"] = {
                        "value": f"{gm:.1f}%", "unit": "%"
                    }

            # Derived EBITDA Margin from ebitda + revenue if not stored
            if "ebitda_margin" not in fm.get("profit_margins", {}):
                rev    = latest.get((company_id, "revenue"))
                ebitda = latest.get((company_id, "ebitda"))
                if rev and rev != 0 and ebitda is not None:
                    em = ebitda / rev * 100
                    fm.setdefault("profit_margins", {})["ebitda_margin"] = {
                        "value": f"{em:.1f}%", "unit": "%"
                    }

            results.append({
                "id":   f"legacy_{company_id}_{period_id}",
                "date": period_id,
                "data": {
                    "financial_metrics_2025": fm,
                    "_source":    "bigquery_legacy",
                    "_period_id": period_id,
                },
                "metadata": {
                    "original_filename": f"histórico {period_id}",
                    "founder_email":     "",
                    "file_hash":         "",
                    "processed_at":      period_id,
                    "gcs_path":          "",
                    "company_domain":    company_id,
                    "portfolio_id":      lookup_portfolio(company_id),
                },
            })

        print(f"[API/all/BQ] {len(results)} empresas históricas desde BigQuery")
        return results

    except Exception as _err:
        print(f"[API/all/BQ] Fallback failed (non-fatal): {_err}")
        return []


def _query_rag_context(portfolio_id: str | None, company_id: str | None) -> list[dict]:
    """
    Fetches the latest valid KPI rows from BigQuery to use as RAG context.
    Returns up to 400 rows ordered by submission date DESC.
    """
    ds = f"{PROJECT_ID}.{os.getenv('BIGQUERY_DATASET', 'cometa_vault')}"
    filters = ["f.is_valid = TRUE", "f.raw_value IS NOT NULL"]
    params  = []

    from google.cloud import bigquery as bq

    if portfolio_id:
        filters.append("s.portfolio_id = @portfolio_id")
        params.append(bq.ScalarQueryParameter("portfolio_id", "STRING", portfolio_id))
    if company_id:
        filters.append("LOWER(s.company_id) LIKE @company_id")
        params.append(bq.ScalarQueryParameter("company_id", "STRING", f"%{company_id.lower()}%"))

    where = " AND ".join(filters)
    sql = f"""
        SELECT
            s.company_id,
            s.portfolio_id,
            s.period_id,
            f.kpi_label,
            f.raw_value,
            f.unit,
            COALESCE(f.is_manually_edited, FALSE) AS is_manually_edited
        FROM `{ds}.fact_kpi_values` f
        JOIN `{ds}.submissions`      s ON f.submission_id = s.submission_id
        WHERE {where}
        ORDER BY s.submitted_at DESC
        LIMIT 400
    """
    try:
        client = _get_bq_client_for_api()
        job    = client.query(sql, job_config=bq.QueryJobConfig(query_parameters=params))
        rows   = list(job.result())
        return [dict(r) for r in rows]
    except Exception as e:
        print(f"[RAG] BQ query failed: {e}")
        return []


# ── KPI Dictionary for RAG ────────────────────────────────────────────────────

def _fetch_kpi_dict_for_rag() -> dict[str, dict]:
    """
    Fetch the full KPI dictionary from dim_kpi_metadata keyed by kpi_key.

    Returns a dict mapping kpi_key → {display_name, description, unit,
    min_historical_year, vertical} for use in the Gemini prompt.

    Non-fatal: on BQ error returns an empty dict so the RAG prompt is
    built without metadata (graceful degradation, no 500 thrown).
    """
    try:
        items = query_kpi_metadata(vertical=None)
        return {
            item["kpi_key"]: {
                "display_name":        item.get("display_name", ""),
                "description":         item.get("description", ""),
                "unit":                item.get("unit", ""),
                "min_historical_year": item.get("min_historical_year"),
                "vertical":            item.get("vertical", "GENERAL"),
            }
            for item in items
        }
    except Exception as e:
        print(f"⚠️  [RAG/dict] KPI metadata fetch failed (non-fatal): {e}")
        return {}


# ── A3: RAG Leak Protection ───────────────────────────────────────────────────

def _verify_rag_integrity(
    rows: list[dict],
    expected_company_id: str,
) -> list[dict]:
    """
    Post-fetch verification that every BQ row belongs to the requested company.

    Controle A3: La query de BigQuery ya filtra por company_id, pero esta función
    actúa como segunda línea de defensa en caso de que el filtro LIKE sea demasiado
    permisivo o sea bypasseado por una condición de carrera.

    Lógica:
    - Si expected_company_id está vacío, no hay restricción → devuelve todo.
    - Para cada row, verifica que su company_id CONTENGA la cadena esperada
      (case-insensitive). Rows que no coincidan son "contaminados".
    - Rows contaminados se eliminan del contexto y se emite una SECURITY ALERT.
    - Si TODOS los rows son contaminados (bypass total del filtro BQ),
      se lanza HTTPException 500 — generación abortada.

    Returns
    -------
    list[dict]  — Solo rows validados para el company_id solicitado.

    Raises
    ------
    HTTPException 500  — Si la contaminación es total (fuga de datos detectada).
    """
    if not expected_company_id or not rows:
        return rows

    needle = expected_company_id.lower().strip()
    clean: list[dict] = []
    contaminated: list[dict] = []

    for row in rows:
        row_company = str(row.get("company_id") or "").lower().strip()
        # Bidirectional containment: handles "solvento" ↔ "solvento.com"
        if needle in row_company or row_company in needle:
            clean.append(row)
        else:
            contaminated.append(row)

    if contaminated:
        leaked_companies = list({r.get("company_id", "?") for r in contaminated})
        print(
            f"🚨 [RAG/A3] SECURITY ALERT — {len(contaminated)} row(s) contaminado(s) "
            f"de {len(rows)} para company='{expected_company_id}'. "
            f"Companies ajenas detectadas: {leaked_companies}"
        )
        if not clean:
            # Contaminación total — el filtro BQ puede haber sido bypasseado.
            raise HTTPException(
                status_code=500,
                detail=(
                    "[A3] Integridad del contexto comprometida: ningún dato "
                    "pertenece a la empresa solicitada. Consulta abortada por seguridad."
                ),
            )

    return clean


# ── Multi-format processing helpers ──────────────────────────────────────────

# Maximum rows rendered per sheet to stay within Gemini's token budget.
# 500 rows ≈ 40–80 KB of Markdown — well within the 1 M-token context window.
_MAX_ROWS_PER_SHEET = 500


def _df_to_markdown(df: pd.DataFrame) -> str:
    """
    Convert a DataFrame to a GitHub-flavored Markdown table.
    No external dependencies (no tabulate required).
    """
    if df.empty:
        return "*(tabla vacía)*"

    # Stringify column names
    cols = [str(c).strip() for c in df.columns]
    header    = "| " + " | ".join(cols) + " |"
    separator = "| " + " | ".join(["---"] * len(cols)) + " |"

    rows = []
    for _, row in df.iterrows():
        cells = [
            str(v).replace("|", "\\|").strip() if pd.notna(v) else ""
            for v in row
        ]
        rows.append("| " + " | ".join(cells) + " |")

    return "\n".join([header, separator] + rows)


def _process_tabular(file_path: str, ext: str, gemini, prompt_schema: str) -> str:
    """
    Read ALL sheets from CSV / XLSX / PARQUET, convert each to Markdown,
    then call Gemini with the identical financial-audit prompt used for PDFs.

    Changes vs. previous implementation:
    - XLSX: reads every sheet (sheet_name=None), not just the first one.
    - Row cap raised to _MAX_ROWS_PER_SHEET (500) per sheet.
    - Output format: GitHub-flavored Markdown tables (better structure preservation).
    - Prompt: uses the full FASE 1 + FASE 2 schema (same as PDF pipeline).
    """
    print(f"📊 [Tabular] Leyendo archivo {ext} (todas las hojas)...")

    try:
        sheets: dict = {}

        if ext == ".csv":
            df = pd.read_csv(file_path, nrows=_MAX_ROWS_PER_SHEET)
            sheets["Hoja1"] = df

        elif ext in (".xlsx", ".xls"):
            # sheet_name=None → devuelve un dict {nombre: DataFrame}
            all_sheets = pd.read_excel(
                file_path, sheet_name=None, engine="openpyxl"
            )
            for name, df in all_sheets.items():
                sheets[str(name)] = df.head(_MAX_ROWS_PER_SHEET)

        elif ext == ".parquet":
            df = pd.read_parquet(file_path).head(_MAX_ROWS_PER_SHEET)
            sheets["Hoja1"] = df

        else:
            raise ValueError(f"Extensión tabular no reconocida: {ext}")

    except Exception as e:
        print(f"❌ [Tabular] Error leyendo archivo: {e}")
        raise RuntimeError(f"No se pudo leer el archivo {ext}: {e}") from e

    _sheet_keys = list(sheets.keys())
    print(f"   Pestañas encontradas: {_sheet_keys}")

    # ── Build Markdown content (one section per sheet) ────────────────────
    md_sections = []
    for sheet_name, df in sheets.items():
        # Drop columns that are 100 % null — they add noise, no signal
        df = df.dropna(axis=1, how="all")
        _n_rows = df.shape[0]
        _n_cols = df.shape[1]
        print(f"   Hoja '{sheet_name}': {_n_rows} filas × {_n_cols} columnas")
        md = _df_to_markdown(df)
        md_sections.append(f"## Pestaña: {sheet_name}\n\n{md}")

    full_markdown = "\n\n---\n\n".join(md_sections)
    _md_len      = len(full_markdown)
    _sheet_count = len(sheets)
    print(f"   Markdown generado: {_md_len:,} caracteres, {_sheet_count} pestaña(s)")

    # ── Prompt: same FASE 1 + FASE 2 schema as the PDF pipeline ─────────
    # Prepend a short adapter note so Gemini knows the source is a spreadsheet,
    # then inject the full audit schema unchanged.
    _ext_upper   = ext.upper()
    _n_sheets    = len(sheets)
    adapter_header = (
        f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
        f"Recibes el contenido completo de un archivo {_ext_upper} con "
        f"{_n_sheets} pestaña(s), convertido a tablas Markdown.\n"
        f"Analiza TODAS las pestañas para localizar las métricas financieras.\n"
        f"Cuando una métrica aparezca en varias hojas, usa la fuente más reciente o la más detallada.\n"
        f"Si los valores están en términos absolutos y tienes los ingresos, calcula los márgenes (%).\n"
        f"Para Revenue Growth YoY: (valor_último - valor_anterior) / |valor_anterior| × 100.\n\n"
        f"Aplica las siguientes instrucciones de extracción EXACTAMENTE:\n\n"
    )
    full_prompt = adapter_header + prompt_schema

    # Pass the prompt as first arg and the Markdown content as second —
    # analizar_texto calls generate_content([prompt, contenido_texto]).
    return gemini.analizar_texto(full_prompt, full_markdown)


def _process_docx(file_path: str, gemini, prompt_schema: str) -> str:
    """Extract text from DOCX and send to Gemini with the financial audit prompt."""
    print(f"📝 [DOCX] Extrayendo texto...")
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)
        full_text = "\n".join(paragraphs)
        if table_texts:
            full_text += "\n\nTABLAS DEL DOCUMENTO:\n" + "\n".join(table_texts)
    except ImportError:
        raise RuntimeError("python-docx no está instalado. Ejecuta: pip install python-docx")
    except Exception as e:
        print(f"❌ [DOCX] Error extrayendo texto: {e}")
        raise RuntimeError(f"No se pudo procesar el DOCX: {e}") from e

    print(f"   Texto extraído: {len(full_text)} caracteres, {len(paragraphs)} párrafos")

    docx_instruction = (
        "Eres un auditor financiero senior. Analiza el siguiente documento Word (DOCX) "
        "que contiene información financiera de una startup. Extrae las métricas del texto "
        "y tablas, y emite el JSON estándar.\n\nCONTENIDO DEL DOCUMENTO:\n"
    )
    return gemini.analizar_texto(docx_instruction + full_text + "\n\n" + prompt_schema, "")


# ── PDF helpers ───────────────────────────────────────────────────────────────

_PDF_CHUNK_SIZE = 90  # páginas máximas por llamada a Gemini

# Sections that MUST exist in the merged Gemini JSON for build_contract() to
# find all 16 KPIs.  If a section is absent (e.g. all chunks had null values)
# we guarantee an empty dict so downstream code never KeyErrors.
_REQUIRED_FM_SECTIONS = (
    "revenue_growth",
    "profit_margins",
    "cash_flow_indicators",
    "debt_ratios",
    "base_metrics",
    "sector_metrics",
)


def _ensure_dict(obj) -> dict:
    """
    Garantiza que el resultado recuperado de GCS sea un dict de Python.

    GCS puede almacenar JSONs con doble serialización (el string ya fue
    serializado una vez antes de ser guardado como string). Este helper
    maneja ambos casos sin excepciones silenciosas.

    Raises TypeError si el objeto no puede convertirse a dict.
    """
    if isinstance(obj, dict):
        return obj
    if isinstance(obj, str):
        parsed = json.loads(obj)
        # Doble serialización: json.loads devolvió otro string
        if isinstance(parsed, str):
            return json.loads(parsed)
        if isinstance(parsed, dict):
            return parsed
        raise TypeError(f"_ensure_dict: json.loads devolvió {type(parsed).__name__}, no dict")
    raise TypeError(f"_ensure_dict: esperaba dict o str, recibió {type(obj).__name__}")


def _ensure_fm_sections(gemini_json: dict) -> dict:
    """
    Garantiza que financial_metrics_2025 tenga todas las sub-secciones
    requeridas para que build_contract() pueda iterar KPI_REGISTRY completo.

    Modifica el dict en-lugar y también lo devuelve (para encadenamiento).
    """
    fm = gemini_json.setdefault("financial_metrics_2025", {})
    for section in _REQUIRED_FM_SECTIONS:
        fm.setdefault(section, {})
    return gemini_json


def split_pdf_to_chunks(file_bytes: bytes, size: int = _PDF_CHUNK_SIZE) -> list[bytes]:
    """
    Divide los bytes de un PDF en fragmentos de máximo `size` páginas.

    Devuelve una lista de bytes — cada elemento es un PDF autónomo válido
    que puede guardarse en disco o enviarse directamente a Gemini.

    Parameters
    ----------
    file_bytes : bytes — contenido completo del PDF original.
    size       : int   — máximo de páginas por fragmento (default 50).

    Returns
    -------
    list[bytes] con 1..N fragmentos. Si el PDF tiene ≤ size páginas,
    la lista tiene exactamente un elemento (los bytes originales sin modificar).
    """
    try:
        import fitz
    except ImportError:
        raise RuntimeError("PyMuPDF no instalado. Ejecuta: pip install pymupdf")

    src_doc     = fitz.open(stream=file_bytes, filetype="pdf")
    total_pages = len(src_doc)

    if total_pages <= size:
        src_doc.close()
        return [file_bytes]

    n_chunks = (total_pages + size - 1) // size
    chunks: list[bytes] = []

    for idx in range(n_chunks):
        start = idx * size
        end   = min(start + size - 1, total_pages - 1)

        chunk_doc = fitz.open()
        chunk_doc.insert_pdf(src_doc, from_page=start, to_page=end)
        chunks.append(chunk_doc.tobytes())
        chunk_doc.close()

        _blk_num    = idx + 1
        _page_start = start + 1
        _page_end   = end + 1
        print(f"[Chunking] Bloque {_blk_num}/{n_chunks}: páginas {_page_start}–{_page_end}")

    src_doc.close()
    return chunks


def merge_consolidated_results(jsons: list[dict]) -> dict:
    """
    Une los resultados de múltiples llamadas a Gemini (chunks de PDF)
    en un único diccionario consolidado.

    Regla de consolidación por KPI:
    - Si un KPI aparece en varios chunks con value != null, se queda el de
      mayor confidence score.
    - Empate de confidence → gana el chunk anterior (índice más bajo),
      preservando el orden documental del PDF.
    - _document_context se toma siempre del primer chunk que lo contenga.
    - base_metrics y sector_metrics se garantizan presentes en el resultado
      (secciones vacías si ningún chunk extrajo datos para ellas).

    Parameters
    ----------
    jsons : list de dicts Gemini ya parseados — mínimo uno.

    Returns
    -------
    Un único dict Gemini con la misma estructura del schema Cometa-Vault.
    """
    import copy

    if not jsons:
        raise ValueError("merge_consolidated_results: lista vacía")
    if len(jsons) == 1:
        return _ensure_fm_sections(jsons[0])

    merged = copy.deepcopy(jsons[0])

    def _get(obj: dict, path: list[str]):
        cur = obj
        for k in path:
            if not isinstance(cur, dict):
                return None
            cur = cur.get(k)
        return cur

    def _set(obj: dict, path: list[str], value) -> None:
        cur = obj
        for k in path[:-1]:
            cur = cur.setdefault(k, {})
        cur[path[-1]] = value

    for chunk_json in jsons[1:]:
        for kpi_def in KPI_REGISTRY:
            path     = kpi_def["path"]
            existing = _get(merged, path)
            incoming = _get(chunk_json, path)

            if not isinstance(incoming, dict) or incoming.get("value") is None:
                continue  # este chunk no tiene dato para este KPI

            if not isinstance(existing, dict) or existing.get("value") is None:
                _set(merged, path, incoming)
                continue

            # Ambos tienen valor — queda el de mayor confidence
            existing_conf = float(existing.get("confidence") or 0.0)
            incoming_conf = float(incoming.get("confidence") or 0.0)
            if incoming_conf > existing_conf:
                _set(merged, path, incoming)

    # Garantizar secciones requeridas aunque todos los chunks tengan null
    _ensure_fm_sections(merged)
    return merged


# Alias de compatibilidad (el upload path lo llamaba merge_kpi_results en
# versiones anteriores — conservamos el nombre por si algún test lo importa)
merge_kpi_results = merge_consolidated_results


def _chunk_and_process_pdf(temp_path: str, gemini, prompt_config: str) -> str:
    """
    Motor de Chunking — lee el PDF desde disco, lo divide en bloques de 50
    páginas con split_pdf_to_chunks(), llama a Gemini por cada bloque y
    consolida con merge_consolidated_results().

    - PDFs ≤ 50 páginas: una sola llamada directa a Gemini (sin overhead).
    - PDFs > 50 páginas: N bloques secuenciales; si un bloque falla se omite
      y se continúa. Si TODOS fallan → RuntimeError.

    Returns
    -------
    String JSON listo para json.loads() — mismo contrato que
    gemini.extraer_y_auditar().
    """
    import re as _re

    with open(temp_path, "rb") as fh:
        file_bytes = fh.read()

    chunks = split_pdf_to_chunks(file_bytes, size=_PDF_CHUNK_SIZE)
    n_chunks = len(chunks)

    if n_chunks == 1:
        print(f"[Chunking] PDF ≤ {_PDF_CHUNK_SIZE} páginas — llamada directa a Gemini")
        return gemini.extraer_y_auditar(temp_path, prompt_config)

    print(f"[Chunking] {n_chunks} bloques de hasta {_PDF_CHUNK_SIZE} páginas")

    chunk_results: list[dict] = []
    for i, chunk_bytes in enumerate(chunks):
        chunk_path = f"{temp_path}_chunk{i}.pdf"
        try:
            with open(chunk_path, "wb") as cf:
                cf.write(chunk_bytes)

            _blk = i + 1
            print(f"[Chunking] Enviando bloque {_blk}/{n_chunks} a Gemini...")
            raw = gemini.extraer_y_auditar(chunk_path, prompt_config)

            if isinstance(raw, str):
                try:
                    parsed = json.loads(raw)
                except json.JSONDecodeError:
                    clean  = _re.sub(r'^```json\s*|\s*```$', '', raw.strip())
                    parsed = json.loads(clean)
            else:
                parsed = raw

            chunk_results.append(parsed)
            print(f"[Chunking] Bloque {_blk} OK")

        except Exception as chunk_err:
            _blk = i + 1
            print(f"[Chunking] Bloque {_blk} falló ({chunk_err}) — omitido")
        finally:
            if os.path.exists(chunk_path):
                os.remove(chunk_path)

    if not chunk_results:
        raise RuntimeError("[Chunking] Ningún bloque fue procesado por Gemini")

    merged = merge_consolidated_results(chunk_results)
    print(f"[Chunking] {len(chunk_results)}/{n_chunks} bloques mergeados exitosamente")
    return json.dumps(merged, ensure_ascii=False)


def check_hash_exists_in_gcs(bucket_name: str, file_hash: str) -> bool:
    """Verifica si existe un archivo con el mismo hash en custom_metadata"""
    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(bucket_name)
        
        # Listar blobs y buscar por metadata
        blobs = bucket.list_blobs()
        
        for blob in blobs:
            if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                print(f"📋 [API] Hash duplicado encontrado en GCS: {file_hash}")
                return True
        
        return False
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f"❌ [API] Error de credenciales/permisos en GCS: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise RuntimeError("GCS_AUTH") from e
    except Exception as e:
        print(f"❌ [API] Error verificando hash en GCS: {e}")
        raise RuntimeError("GCS_ERROR") from e

def get_existing_result(bucket_name: str, file_hash: str) -> dict:
    """Obtiene el resultado JSON existente para un hash específico"""
    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(bucket_name)
        
        # Buscar en staging por hash
        blobs = bucket.list_blobs(prefix="staging/")
        
        for blob in blobs:
            if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                # Descargar y retornar el resultado existente
                content = blob.download_as_text()
                result = json.loads(content)
                print(f"📋 [API] Resultado existente encontrado para hash: {file_hash}")
                return result
        
        return None
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f"❌ [API] Error de credenciales/permisos obteniendo resultado: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise RuntimeError("GCS_AUTH") from e
    except Exception as e:
        print(f"❌ [API] Error obteniendo resultado existente: {e}")
        raise RuntimeError("GCS_ERROR") from e

@app.post("/upload")
@limiter.limit("20/minute")
async def upload_pdf(
    request: Request,
    file: UploadFile = File(...),
    founder_email: str = Header(None, description="Email del founder para identificación"),
    company_id: str = Header(None, description="Company ID para multi-tenancy"),
    token: dict = Depends(_require_auth),
):
    """
    Endpoint para subir PDF y procesar con Gemini (asíncrono)
    """
    # ── DEBUG: log inmediato al primer byte de la request ─────────────────
    print("=" * 60)
    print(f"📥 [DEBUG] /upload HIT — conexión recibida")
    _filename     = getattr(file, "filename", "N/A")
    _content_type = getattr(file, "content_type", "N/A")
    print(f"   filename     : {_filename}")
    print(f"   content_type : {_content_type}")
    print(f"   founder_email: {founder_email!r}")
    print(f"   company_id   : {company_id!r}")
    print("=" * 60)
    # ─────────────────────────────────────────────────────────────────────

    try:
        # ── C5: Validar y sanitizar headers de entrada ──────────────────────
        founder_email = _validate_email_header(founder_email)
        company_id    = _validate_company_header(company_id)

        # 1. Validar extensión
        ALLOWED_EXTENSIONS = {'.pdf', '.csv', '.xlsx', '.xls', '.parquet', '.docx', '.doc'}
        file_ext = os.path.splitext(file.filename or "")[1].lower()
        if not file.filename or file_ext not in ALLOWED_EXTENSIONS:
            _allowed_str = ", ".join(sorted(ALLOWED_EXTENSIONS))
            raise HTTPException(
                status_code=400,
                detail=f"Formato no soportado. Permitidos: {_allowed_str}"
            )
        print(f"📁 [DEBUG] Extensión detectada: {file_ext}")

        # 4. Leer contenido (adelantado para C2 y C7 antes de procesar)
        file_content = await file.read()

        # ── C2: Límite de tamaño ────────────────────────────────────────────
        if len(file_content) > _MAX_FILE_BYTES:
            raise HTTPException(
                status_code=413,
                detail=f"Archivo supera el límite de {_MAX_FILE_MB} MB"
            )

        # ── C7: Validar magic bytes ─────────────────────────────────────────
        if not _validate_magic_bytes(file_content, file_ext):
            raise HTTPException(
                status_code=415,
                detail=f"El contenido binario no corresponde a un archivo {file_ext} válido"
            )

        # 2. Extraer company_id del header o del email
        company_domain = company_id if company_id else (founder_email.split('@')[-1] if founder_email and '@' in founder_email else 'unknown')
        print(f"🏢 [API] Company domain: {company_domain}")
        
        # 3. Si no hay company_id, Vertex AI lo identificará del contenido del PDF
        if not company_domain or company_domain == 'unknown':
            print(f"[API] No company_id en header — se identificará desde el PDF")
            company_domain = 'pending_detection'
        
        # 4. Calcular hash (file_content ya fue leído antes de las validaciones C2/C7)
        file_hash = get_file_hash(file_content)
        
        print(f"📤 [API] Archivo recibido: {file.filename}")
        print(f"🔍 [API] Hash calculado: {file_hash}")
        print(f"👤 [API] Founder: {founder_email}")
        print(f"🏢 [API] Vault path: vault/{company_domain}/")
        
        # 5. Verificar si ya existe por hash en la bóveda específica
        try:
            # Primero verificar en la bóveda de la empresa
            vault_prefix = f"vault/{company_domain}/"
            storage_client = _get_storage_client()
            bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)
            
            # Buscar en la bóveda específica de la empresa
            blobs = bucket.list_blobs(prefix=vault_prefix)
            
            for blob in blobs:
                if blob.metadata and blob.metadata.get('file_hash') == file_hash:
                    # Encontrar resultado existente en la bóveda de la empresa
                    content = blob.download_as_text()
                    # Fix-1: garantizar dict (GCS puede devolver doble-serialización)
                    result = _ensure_dict(json.loads(content))
                    print(f"📋 [API] Resultado duplicado encontrado en bóveda de {company_domain}: {file_hash}")
                    # OBS-04: recalcular checklist_status desde el JSON cacheado
                    try:
                        _dup_contract = build_contract(
                            gemini_json=result,
                            file_hash=file_hash,
                            company_id=company_domain,
                            founder_email=founder_email or "",
                            original_filename=file.filename,
                        )
                        _dup_bucket = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
                        dup_checklist = build_checklist_status(_dup_contract["kpi_rows"], _dup_bucket)
                    except Exception as _ce:
                        print(f"[API] checklist recalc failed for duplicate ({_ce}) — omitting")
                        dup_checklist = None
                    return JSONResponse(
                        content={
                            "status": "success",
                            "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                            "duplicate": True,
                            "result": result,
                            "file_hash": file_hash,
                            "company_domain": company_domain,
                            "checklist_status": dup_checklist,
                        },
                        status_code=200
                    )
            
            # Si no existe en la bóveda de la empresa, buscar en el bucket general
            if check_hash_exists_in_gcs(GCS_INPUT_BUCKET, file_hash):
                # Copiar resultado a la bóveda de la empresa
                existing_result_raw = get_existing_result(GCS_OUTPUT_BUCKET, file_hash)
                if existing_result_raw:
                    # Fix-1: garantizar dict antes de operar con él
                    existing_result = _ensure_dict(existing_result_raw)
                    # Copiar a la bóveda específica
                    vault_result_filename = f"{vault_prefix}{file_hash}_result.json"
                    vault_blob = bucket.blob(vault_result_filename)

                    vault_blob.metadata = {
                        'file_hash': file_hash,
                        'original_filename': existing_result.get('original_filename', 'unknown'),
                        'founder_email': founder_email,
                        'company_domain': company_domain,
                        'vault_path': vault_prefix,
                        'processed_at': datetime.now(timezone.utc).isoformat(),
                        'copied_from_general': True
                    }
                    
                    vault_blob.upload_from_string(
                        json.dumps(existing_result, indent=2),
                        content_type='application/json'
                    )
                    
                    print(f"📋 [API] Resultado copiado a bóveda de {company_domain}: {vault_result_filename}")
                    # OBS-04: recalcular checklist_status desde el JSON cacheado
                    try:
                        _dup_contract2 = build_contract(
                            gemini_json=existing_result,
                            file_hash=file_hash,
                            company_id=company_domain,
                            founder_email=founder_email or "",
                            original_filename=file.filename,
                        )
                        _dup_bucket2 = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
                        dup_checklist2 = build_checklist_status(_dup_contract2["kpi_rows"], _dup_bucket2)
                    except Exception as _ce2:
                        print(f"[API] checklist recalc failed for duplicate-copy ({_ce2}) — omitting")
                        dup_checklist2 = None
                    return JSONResponse(
                        content={
                            "status": "success",
                            "message": "Documento reconocido en la bóveda de Cometa. Sincronizando métricas...",
                            "duplicate": True,
                            "result": existing_result,
                            "file_hash": file_hash,
                            "company_domain": company_domain,
                            "checklist_status": dup_checklist2,
                        },
                        status_code=200
                    )
        except RuntimeError as e:
            if str(e) == "GCS_AUTH":
                raise HTTPException(
                    status_code=500,
                    detail=(
                        "Error de autenticación/permisos con GCS. "
                        "Verifica GOOGLE_APPLICATION_CREDENTIALS y permisos del bucket."
                    ),
                )
            raise HTTPException(
                status_code=500,
                detail="Error de conexión/lectura con GCS durante deduplicación",
            )
        
        # 6. Si es nuevo, iniciar procesamiento asíncrono
        print(f"🆕 [API] Archivo nuevo ({file_ext}), iniciando procesamiento asíncrono...")

        # Guardar temporalmente para procesamiento
        safe_filename = _sanitize_filename(file.filename)  # C6: sanitización segura
        temp_path = os.path.join('/tmp', f"{file_hash}_{safe_filename}")
        os.makedirs('/tmp', exist_ok=True)

        with open(temp_path, "wb") as temp_file:
            temp_file.write(file_content)

        # 7. Iniciar procesamiento con Vertex AI
        try:
            # Inicializar adaptadores
            doc_ai = DocumentAIAdapter(PROJECT_ID, LOCATION_DOC_AI, PROCESSOR_ID)
            print(f"DEBUG: El valor de VERTEX_LOCATION es: '{VERTEX_LOCATION}'")
            gemini = GeminiAuditor(PROJECT_ID, VERTEX_LOCATION)
            
            # ── Contexto de vertical (inyectado antes del prompt) ─────────────
            # El bucket_id se conoce en cuanto detectamos la empresa; pasarlo
            # al prompt evita que Gemini busque métricas de otras verticales.
            _bucket_id = COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            _sector_hints: dict[str, str] = {
                "SAAS":  "Prioriza MRR, Churn Rate y CAC. Si el documento menciona 'Monthly Recurring Revenue', 'Churn' o 'Customer Acquisition Cost', extráelos en sector_metrics.",
                "LEND":  "Prioriza Portfolio Size y NPL Ratio. Si el documento menciona 'cartera de crédito', 'morosidad', 'NPL' o 'Non-Performing Loans', extráelos en sector_metrics.",
                "ECOM":  "Prioriza GMV. Si el documento menciona 'Gross Merchandise Value', 'Total Sales Volume', 'Total Transaction Value' o 'GMV', extráelo en sector_metrics.gmv.",
                "INSUR": "Prioriza Loss Ratio. Si el documento menciona 'siniestralidad', 'claims ratio' o 'loss ratio', extráelo en sector_metrics.",
                "OTH":   "Extrae las métricas financieras estándar. No hay métricas sectoriales obligatorias.",
            }
            _sector_instruction = _sector_hints.get(
                _bucket_id,
                "Extrae todas las métricas del esquema. No hay contexto sectorial específico disponible.",
            )

            # ── Prompt CoT + esquema con confidence ──────────────────────────
            # El prompt se construye en dos partes para evitar conflictos de
            # escape en f-strings: el prefijo dinámico (f-string con bucket_id
            # y sector_instruction) se concatena con el cuerpo estático que
            # contiene el esquema JSON (curly braces literales).
            _prompt_prefix = (
                f"Eres un auditor financiero senior especializado en due diligence de startups.\n"
                f"Tu misión es extraer métricas financieras de un PDF con precisión institucional.\n"
                f"\n"
                f"╔══════════════════════════════════════════════════════════════╗\n"
                f"║  CONTEXTO DE INDUSTRIA                                      ║\n"
                f"╚══════════════════════════════════════════════════════════════╝\n"
                f"\n"
                f"Estás analizando una empresa de la vertical {_bucket_id}.\n"
                f"{_sector_instruction}\n"
                f"No intentes inventar o extraer métricas de otras verticales a menos que el\n"
                f"documento las mencione explícitamente como KPIs de negocio de la empresa.\n"
                f"\n"
            )
            _prompt_body = """
╔══════════════════════════════════════════════════════════════╗
║  FASE 1 — ANÁLISIS PREVIO  (escribe en _document_context)  ║
╚══════════════════════════════════════════════════════════════╝

Ejecuta este análisis y materializa las conclusiones en el campo
`_document_context` del JSON de salida (ver esquema más abajo).

  A. MONEDA PRINCIPAL
     ¿Cuál es la moneda del documento? (USD, MXN, EUR, COP, etc.)
     Si hay más de una moneda, ¿cuál domina los estados financieros?

  B. PERÍODO DE REPORTE
     ¿Qué año fiscal o período cubre el documento?
     Busca encabezados como "FY2025", "Año terminado el 31/12/2025",
     "H1 2025", "Q4 2025". Anota si el documento cubre períodos parciales.

  C. ESCALA NUMÉRICA
     ¿Los montos están expresados en unidades base, miles ($K) o millones ($M)?
     Busca notas al pie como "en miles de pesos" o "amounts in USD thousands".
     Normaliza TODOS los valores al output usando K, M o B según corresponda.

  D. ZONAS DE AMBIGÜEDAD
     Identifica métricas donde:
       · El dato no aparece explícito y debes calcularlo o inferirlo.
       · Hay dos cifras posibles (ej. consolidado vs. entidad separada).
       · El documento es un deck de presentación sin estados auditados.
     Para cada zona de ambigüedad, prepara una explicación para el campo
     "description" y asigna confidence < 0.70.

Las conclusiones de A, B y C van en `_document_context`.
Las zonas de D informan los campos description y confidence de cada métrica.

╔══════════════════════════════════════════════════════════════╗
║  FASE 2 — EXTRACCIÓN JSON                                   ║
╚══════════════════════════════════════════════════════════════╝

REGLAS OBLIGATORIAS:
1. Responde ÚNICAMENTE con el objeto JSON. Cero caracteres fuera del JSON.
2. Usa EXACTAMENTE las claves del esquema. Sin sinónimos, sin claves extra.
3. Cada métrica contiene tres campos obligatorios:

   "value"
     El número con unidad y escala normalizada según FASE 1.
     Ejemplos: "36%", "$9.7M", "-$320K", "-0.74%".
     NUNCA omitas la unidad. NUNCA escribas un número puro.

   "confidence"
     Float 0.0–1.0 que refleja tu certeza sobre el valor extraído:
       >= 0.90  → dato explícito, sin ambigüedad, fuente directa.
       0.70–0.89 → dato requirió cálculo menor o inferencia razonable.
       < 0.70  → dato ambiguo, estimado, parcial o de fuente indirecta.
     Sé honesto: subestimar confidence es mejor que inflar confianza falsa.

   "description"
     Cita exacta de la fuente: tabla, línea y página/sección del documento.
     Si confidence < 0.70: DEBE incluir (a) por qué el dato es incierto,
     (b) qué alternativas se consideraron y (c) cuál fue el criterio de
     selección del valor reportado.

4. Si una métrica no aparece en el documento: escribe null (no "N/A", no "---").
5. La clave raíz SIEMPRE debe ser "financial_metrics_2025".
6. Normaliza la escala usando el contexto de FASE 1 antes de escribir cada value.
7. El campo `_document_context` es OBLIGATORIO. Escribe exactamente sus 4 sub-campos.
   currency usa el código ISO 4217 de 3 letras (USD, MXN, EUR, BRL, COP, ARS...).

ESQUEMA REQUERIDO:
{
  "_document_context": {
    "currency":    "<ISO 4217 de la moneda dominante del documento, ej. 'MXN'>",
    "period":      "<período fiscal, ej. 'FY2025', 'H1 2025', 'Q4 2025'>",
    "scale":       "<escala de montos: 'units', 'thousands', 'millions', 'billions'>",
    "scale_notes": "<dónde se encontró la indicación de escala en el doc, o null>"
  },
  "financial_metrics_2025": {
    "revenue_growth": {
      "value": "<crecimiento YoY en %, ej. '36%' o '-4.2%'>",
      "confidence": <float 0.0-1.0>,
      "description": "<tabla/línea de origen; si confidence<0.70: razonamiento>"
    },
    "profit_margins": {
      "gross_profit_margin": {
        "value": "<%, ej. '68%'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      },
      "ebitda_margin": {
        "value": "<%, puede ser negativo, ej. '-12%'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      }
    },
    "cash_flow_indicators": {
      "cash_in_bank_end_of_year": {
        "value": "<monto con símbolo y escala, ej. '$9.7M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      },
      "annual_cash_flow": {
        "value": "<monto, ej. '-$3.2M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      }
    },
    "debt_ratios": {
      "working_capital_debt": {
        "value": "<monto, ej. '$1.1M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      }
    },
    "base_metrics": {
      "revenue": {
        "value": "<monto absoluto con símbolo y escala, ej. '$4.2M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      },
      "ebitda": {
        "value": "<monto, puede ser negativo, ej. '-$0.8M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      },
      "cogs": {
        "value": "<costo de ventas, ej. '$1.3M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente>"
      }
    },
    "sector_metrics": {
      "mrr": {
        "value": "<Monthly Recurring Revenue, ej. '$350K'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — solo para SaaS>"
      },
      "churn_rate": {
        "value": "<%mensual de cancelaciones, ej. '2.1%'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — solo para SaaS>"
      },
      "cac": {
        "value": "<Customer Acquisition Cost, ej. '$120'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — SaaS/ECOM/INSUR>"
      },
      "portfolio_size": {
        "value": "<cartera total de créditos, ej. '$25M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — solo para Lending>"
      },
      "npl_ratio": {
        "value": "<Non-Performing Loan ratio, ej. '3.4%'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — solo para Lending>"
      },
      "gmv": {
        "value": "<Gross Merchandise Value / Total Sales Volume / Total Transaction Value / GMV, ej. '$8.5M'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — eCommerce/logística. Aliases aceptados: 'Total Sales Volume', 'Total Transaction Value', 'GMV'>"
      },
      "loss_ratio": {
        "value": "<siniestralidad, ej. '62%'>",
        "confidence": <float 0.0-1.0>,
        "description": "<fuente — solo para Insurtech>"
      }
    }
  }
}

Analiza el documento adjunto y responde con el JSON completo. Nada más.
"""
            prompt_config = _prompt_prefix + _prompt_body

            # ── Enrutar por tipo de archivo ────────────────────────────────
            if file_ext in ('.csv', '.xlsx', '.xls', '.parquet'):
                resultado_raw = _process_tabular(temp_path, file_ext, gemini, prompt_config)
            elif file_ext in ('.docx', '.doc'):
                resultado_raw = _process_docx(temp_path, gemini, prompt_config)
            else:
                # PDF — Smart Chunking (bloques de 50 páginas)
                resultado_raw = _chunk_and_process_pdf(temp_path, gemini, prompt_config)

            # Normalizar: Gemini devuelve un string JSON; lo parseamos para evitar
            # doble serialización al guardarlo en GCS.
            if isinstance(resultado_raw, str):
                try:
                    resultado = json.loads(resultado_raw)
                except json.JSONDecodeError:
                    # Si el modelo devolvió texto envuelto en ```json ... ```, limpiarlo
                    import re
                    clean = re.sub(r'^```json\s*|\s*```$', '', resultado_raw.strip())
                    resultado = json.loads(clean)
            else:
                resultado = resultado_raw
            
            # 8a. Identificación silenciosa — Vertex AI detecta la empresa del PDF.
            # Escanea el JSON de Gemini + nombre de archivo para identificar la startup.
            # Si se detecta, sobrescribe company_domain y asigna portfolio_id desde dim_company.
            detected_key, detected_portfolio = detect_company_from_text(
                json.dumps(resultado) + " " + file.filename
            )
            if detected_key != "unknown":
                print(f"[API] Empresa auto-detectada: '{detected_key}' -> Fondo {detected_portfolio}")
                company_domain = detected_key
                portfolio_id   = detected_portfolio
            else:
                portfolio_id = lookup_portfolio(company_domain)
                print(f"[API] Empresa no detectada del contenido, usando header: {company_domain} -> {portfolio_id}")

            # 8b. Build the canonical data contract (Rule 4 + Rule 8)
            contract = build_contract(
                gemini_json=resultado,
                file_hash=file_hash,
                company_id=company_domain,
                founder_email=founder_email or "",
                original_filename=file.filename,
                portfolio_id=portfolio_id,
            )

            # ── R1 + R2: Normalización de IDs canónicos ──────────────────────────
            # Muta el contrato in-place ANTES de leer ningún campo:
            #   • period_id  →  PYYYYQxMyy  (ej. "P2025Q1M03")
            #   • company_id →  COMP_XXX    (ej. "COMP_SOLVENTO")
            # También propaga company_id a cada kpi_row (requerido por el contrato JSON).
            _raw_period = resultado.get("_document_context", {}).get("period", "") or ""
            _norm_result = _apply_contract_normalization(
                contract=contract,
                raw_company=company_domain,
                raw_period=_raw_period,
            )
            # Sincronizar company_domain local con la clave canónica lowercase para GCS.
            # IMPORTANTE: el contrato ya tiene comp_id ("COMP_SOLVENTO") mutado para BQ.
            # Las rutas GCS deben usar la clave lowercase ("solvento") para coincidir
            # con lo que /api/portfolio-companies devuelve al sidebar del analista.
            if _norm_result["company_ok"]:
                # Empresa conocida: derivar clave lowercase desde comp_id (strip "COMP_")
                company_domain = _norm_result["comp_id"].replace("COMP_", "").lower()
            else:
                # Empresa desconocida: normalizar dominio (strip TLD, lowercase)
                company_domain = company_domain.lower().split(".")[0].replace("-", "").replace("_", "")
            portfolio_id = _norm_result["fund_id"] or portfolio_id

            print(
                f"🔖 [R1/R2] Normalization — "
                f"period: '{_norm_result['period_id']}' (ok={_norm_result['period_ok']}) | "
                f"company: '{_norm_result['comp_id']}' (known={_norm_result['company_ok']}) | "
                f"fund: '{_norm_result['fund_id']}' | bucket: '{_norm_result['bucket_id']}'"
            )
            if _norm_result["errors"]:
                for _err in _norm_result["errors"]:
                    print(f"⚠️  [R1/R2] {_err}")
            # ─────────────────────────────────────────────────────────────────────

            integrity = contract["integrity"]
            _sub        = contract["submission"]
            _kpi_valid  = _sub["kpi_count_valid"]
            _kpi_total  = _sub["kpi_count_total"]
            _period_id  = _sub["period_id"]   # ya normalizado a PYYYYQxMyy
            _period_ok  = integrity["period_consistent"]
            print(
                f"📋 [API] Contract built — "
                f"valid KPIs: {_kpi_valid}/{_kpi_total}, "
                f"period: {_period_id}, "
                f"period_consistent: {_period_ok}"
            )

            if integrity["warnings"]:
                for w in integrity["warnings"]:
                    print(f"⚠️  [API] Integrity warning: {w}")

            # 8c. Sector checklist — usa bucket_id resuelto por R2
            company_bucket = _norm_result["bucket_id"] or COMPANY_BUCKET.get(company_domain, "UNKNOWN")
            checklist_status = build_checklist_status(contract["kpi_rows"], company_bucket)
            # Enrich checklist with per-KPI confidence scores so the frontend can
            # highlight low-confidence fields before the founder manually corrects them.
            _conf_scores = _extract_kpi_confidence_scores(resultado)
            if _conf_scores:
                checklist_status["confidence_scores"] = _conf_scores
            if not checklist_status["is_complete"]:
                _missing_kpis = checklist_status["missing_critical_kpis"]
                print(f"[API] Checklist incompleto ({company_bucket}): faltan {_missing_kpis}")

            # ── Validación de contenido financiero ──────────────────────────────
            # Si Gemini no extrajo ningún KPI reconocible, el archivo no es un
            # reporte financiero válido: bloquear toda persistencia y avisar al
            # frontend con 422 (el tempfile ya se limpia más abajo en el except).
            if not _is_financial_document(resultado):
                print(
                    f"🚫 [API] Documento rechazado — sin KPIs financieros reconocibles. "
                    f"Hash: {file_hash}, archivo: {file.filename}"
                )
                if os.path.exists(temp_path):
                    os.remove(temp_path)
                raise HTTPException(
                    status_code=422,
                    detail=(
                        "Archivo no reconocido como reporte financiero válido. "
                        "No se han guardado datos."
                    ),
                )
            # ────────────────────────────────────────────────────────────────────

            # 8d. Persist to BigQuery (non-fatal if BQ is down)
            db_result = {"inserted": False, "duplicate": False, "submission_id": None}
            try:
                db_result = insert_contract(contract)
                if db_result["duplicate"]:
                    print(f"[API] BQ dedup — submission already exists for hash {file_hash}")
            except Exception as db_err:
                print(f"[API] BigQuery write failed (non-fatal, GCS copy kept): {db_err}")

            # 8d. Save to GCS vault bajo el company_domain ya detectado
            storage_client = _get_storage_client()
            bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)
            vault_prefix = f"vault/{company_domain}/"
            result_filename = f"{vault_prefix}{file_hash}_result.json"
            blob = bucket.blob(result_filename)
            blob.metadata = {
                'file_hash': file_hash,
                'original_filename': file.filename,
                'founder_email': founder_email,
                'company_domain': company_domain,
                'portfolio_id': portfolio_id,
                'vault_path': vault_prefix,
                'processed_at': datetime.now(timezone.utc).isoformat()
            }
            blob.upload_from_string(
                json.dumps(resultado, indent=2, ensure_ascii=False),
                content_type='application/json'
            )

            print(f"✅ [API] Resultado guardado en GCS: {result_filename}")

            # Guardar archivo original en GCS (vault/{company}/raw/)
            raw_filename = f"vault/{company_domain}/raw/{file_hash}_{safe_filename}"
            raw_blob = bucket.blob(raw_filename)
            raw_blob.metadata = {
                'file_hash': file_hash,
                'original_filename': file.filename,
                'founder_email': founder_email,
                'company_domain': company_domain,
                'file_type': file_ext,
                'uploaded_at': pd.Timestamp.now().isoformat(),
            }
            raw_blob.upload_from_string(file_content, content_type=file.content_type or 'application/octet-stream')
            print(f"📦 [API] Archivo original guardado en GCS: {raw_filename}")

            # 8e. Receipt email — Vault Seal SHA-256 (non-fatal)
            try:
                from src.services.hash_service import generate_vault_seal
                from src.services.email_service import send_receipt_email
                _processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
                _vault_seal = generate_vault_seal(
                    company_id   = company_domain,
                    file_hash    = file_hash,
                    kpi_rows     = contract["kpi_rows"],
                    processed_at = _processed_at,
                )
                send_receipt_email(
                    to_email       = founder_email,
                    company_domain = company_domain,
                    period         = _period_id,
                    vault_seal     = _vault_seal,
                    file_hash      = file_hash,
                    kpi_count      = _kpi_valid,
                    processed_at   = _processed_at,
                )
            except Exception as _receipt_err:
                print(f"[API] Receipt email failed (non-fatal): {_receipt_err}")

            # 9. Limpiar archivo temporal
            os.remove(temp_path)

            # 10. Retornar contrato completo al frontend
            kpi_confidence_scores = _extract_kpi_confidence_scores(resultado)
            return JSONResponse(
                content={
                    "status": "success",
                    "message": "Archivo procesado exitosamente",
                    "duplicate": False,
                    # Legacy field: raw Gemini dict for existing frontend consumers
                    "result": resultado,
                    # Structured contract for The Vault
                    "submission":        contract["submission"],
                    "kpi_rows":          contract["kpi_rows"],
                    "integrity":         integrity,
                    "db":                db_result,
                    "file_hash":         file_hash,
                    "company_domain":    company_domain,
                    # Sector checklist — feedback inmediato al founder
                    "checklist_status":  checklist_status,
                    # Per-KPI confidence scores extracted from Gemini (0–100 integer)
                    "kpi_confidence_scores": kpi_confidence_scores,
                },
                status_code=200
            )
            
        except Exception as e:
            # Limpiar archivo temporal en caso de error
            if os.path.exists(temp_path):
                os.remove(temp_path)
            
            print(f"❌ [API] Error en procesamiento: {str(e)}")
            traceback.print_exc()
            
            # Error claro para el usuario
            return JSONResponse(
                content={
                    "status": "error",
                    "message": "Error procesando el archivo",
                    "error": str(e)
                },
                status_code=500
            )
    
    except Exception as e:
        print(f"❌ [API] Error general en upload: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error en el servidor: {str(e)}")


class ManualUpdateRequest(BaseModel):
    """Body for POST /api/founder/manual-update."""
    file_hash: str
    updates:   dict[str, str]


@app.post("/api/founder/manual-update")
@limiter.limit("30/minute")
async def founder_manual_update(
    request: Request,
    body: ManualUpdateRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Persist founder-supplied corrections for missing or low-confidence KPIs.

    The endpoint locates the stored result JSON in GCS by ``file_hash`` within
    the founder's company vault, applies the key/value corrections supplied in
    ``updates``, and overwrites the blob.  This is a best-effort write: if the
    blob is not found or GCS is unavailable, a descriptive 404/500 is returned.

    Parameters
    ----------
    body.file_hash : str
        SHA-256 prefix (16 chars) that identifies the processed document.
    body.updates   : dict[str, str]
        Mapping of KPI key → corrected value string, e.g.
        ``{"mrr": "$350K", "churn_rate": "2.1%"}``.

    Returns
    -------
    JSON ``{ "status": "ok", "updated_fields": [str] }``
    """
    company_domain: str = token.get("company_id") or token.get("sub", "")
    # Normalise to domain-only (strip full email if necessary)
    if "@" in company_domain:
        company_domain = company_domain.split("@")[-1]

    if not company_domain:
        raise HTTPException(status_code=403, detail="company_id no disponible en el token")

    vault_prefix = f"vault/{company_domain}/"

    try:
        storage_client = _get_storage_client()
        bucket_obj     = storage_client.bucket(GCS_OUTPUT_BUCKET)
        blobs          = list(bucket_obj.list_blobs(prefix=vault_prefix))
    except Exception as gcs_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al conectar con GCS: {gcs_err}",
        )

    # Find the blob for this file_hash
    target_blob = None
    for blob in blobs:
        if blob.name.endswith(".json") and body.file_hash in blob.name:
            target_blob = blob
            break

    if target_blob is None:
        raise HTTPException(
            status_code=404,
            detail=f"No se encontró resultado para el hash '{body.file_hash}' en la bóveda de {company_domain}",
        )

    try:
        existing_raw  = target_blob.download_as_text()
        existing_data = _ensure_dict(json.loads(existing_raw))
    except Exception as read_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al leer el resultado existente: {read_err}",
        )

    # Apply corrections — inject into the manual_corrections sub-object
    manual_corrections: dict = existing_data.get("manual_corrections") or {}
    for k, v in body.updates.items():
        manual_corrections[k] = v
    existing_data["manual_corrections"] = manual_corrections

    try:
        target_blob.upload_from_string(
            json.dumps(existing_data, indent=2, ensure_ascii=False),
            content_type="application/json",
        )
    except Exception as write_err:
        raise HTTPException(
            status_code=500,
            detail=f"Error al guardar correcciones en GCS: {write_err}",
        )

    updated_fields = list(body.updates.keys())
    print(
        f"[manual-update] {len(updated_fields)} campo(s) corregidos para "
        f"hash={body.file_hash} company={company_domain}: {updated_fields}"
    )
    return JSONResponse(
        content={"status": "ok", "updated_fields": updated_fields},
        status_code=200,
    )


@app.get("/api/result/{file_hash}")
async def get_analysis_result(file_hash: str):
    """
    Obtiene el resultado del análisis por hash de archivo
    """
    try:
        result = get_existing_result(GCS_OUTPUT_BUCKET, file_hash)
        if result:
            return JSONResponse(content=result)
        else:
            raise HTTPException(status_code=404, detail="Resultado no encontrado")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error obteniendo resultado: {str(e)}")

@app.get("/api/results")
async def get_all_results(company_id: str = None, token: dict = Depends(_require_auth)):
    """
    Obtiene todos los resultados de análisis guardados en GCS/vault/{company_id}/

    Multi-tenancy rules
    -------------------
    ANALISTA (ANA-*)  — can query any company_id supplied in the URL.
    FOUNDER  (FND-*)  — the URL company_id is ignored; the company is always
                        derived from the email claim in the JWT (Zero Trust).
                        Returns 403 if the token carries no usable company.
    """
    # ── Multi-tenancy gate ─────────────────────────────────────────────────────
    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_founder = (role == "FOUNDER") or user_id.startswith("FND-")
    if is_founder:
        # Derive company from JWT — never trust the URL param for Founders.
        # token.get("company_id") covers tokens minted with an explicit claim;
        # falling back to "sub" (email) and stripping the domain works for
        # standard Founder accounts whose email matches their company domain.
        raw_company: str = (token.get("company_id") or token.get("sub") or "").strip()
        if "@" in raw_company:
            raw_company = raw_company.split("@")[-1]   # john@solvento.com → solvento.com
        jwt_company = raw_company.lower()
        if not jwt_company:
            raise HTTPException(
                status_code=403,
                detail="Founder sin company_id en el token. Contacta a tu analista.",
            )
        # Override whatever the caller sent — scope is enforced server-side
        company_id = jwt_company

    try:
        storage_client = _get_storage_client()
        bucket = storage_client.bucket(GCS_OUTPUT_BUCKET)

        # Si no se proporciona company_id, devolver error
        if not company_id:
            raise HTTPException(status_code=400, detail="company_id es obligatorio")

        # Buscar en la ruta canónica y en la ruta legada COMP_XXX para compatibilidad
        # con documentos subidos antes del fix de normalización de company_domain.
        cid_clean = company_id.lower().strip()
        vault_prefixes = [
            f"vault/{cid_clean}/",                      # ruta canónica: vault/solvento/
            f"vault/COMP_{cid_clean.upper()}/",         # ruta legada:   vault/COMP_SOLVENTO/
        ]

        results = []
        seen_ids: set[str] = set()  # evitar duplicados si por alguna razón coinciden

        for vault_prefix in vault_prefixes:
            for blob in bucket.list_blobs(prefix=vault_prefix):
                if not blob.name.endswith('.json'):
                    continue
                # Deduplicate by the computed id (hash-based), not the full blob path.
                # Two blobs from canonical and legacy prefixes can share the same derived
                # id if the same file exists in both; track that to avoid React key errors.
                blob_id = blob.name.replace('.json', '').replace(vault_prefix, '')
                if blob_id in seen_ids:
                    continue
                seen_ids.add(blob_id)
                try:
                    # Descargar contenido del JSON
                    content = blob.download_as_text()
                    
                    # Manejar posible doble serialización
                    try:
                        result_data = json.loads(content)
                        # Si el resultado sigue siendo un string, aplicar segundo parse
                        if isinstance(result_data, str):
                            print(f" [API] Detectada doble serialización en {blob.name}")
                            result_data = json.loads(result_data)
                    except json.JSONDecodeError as e:
                        print(f" [API] Error parseando JSON de {blob.name}: {e}")
                        continue
                    
                    # Asegurarse que result_data sea un diccionario
                    if not isinstance(result_data, dict):
                        _blob_name      = blob.name
                        _result_type    = type(result_data).__name__
                        print(f" [API] Resultado no es diccionario en {_blob_name}: {_result_type}")
                        continue
                    
                    # Extraer metadata
                    metadata = blob.metadata or {}
                    
                    # Construir objeto de resultado con estructura clara
                    # Derive portfolio — prefer stored metadata, fall back to lookup
                    blob_portfolio = metadata.get('portfolio_id') or lookup_portfolio(
                        metadata.get('company_domain', company_id)
                    )
                    result_item = {
                        "id": blob.name.replace('.json', '').replace(vault_prefix, ''),
                        "data": result_data,
                        "date": metadata.get('processed_at', 'unknown'),
                        "metadata": {
                            "original_filename": metadata.get('original_filename', 'unknown'),
                            "founder_email": metadata.get('founder_email', 'unknown'),
                            "file_hash": metadata.get('file_hash', ''),
                            "processed_at": metadata.get('processed_at', 'unknown'),
                            "gcs_path": blob.name,
                            "company_domain": metadata.get('company_domain', company_id),
                            "portfolio_id": blob_portfolio,
                        }
                    }

                    results.append(result_item)
                    _blob_name   = blob.name
                    _data_type   = type(result_data).__name__
                    _data_keys   = list(result_data.keys()) if isinstance(result_data, dict) else "N/A"
                    print(f" [API] Resultado cargado: {_blob_name}")
                    print(f" [API] Tipo de data: {_data_type}")
                    print(f" [API] Keys en data: {_data_keys}")
                    
                except Exception as e:
                    print(f" [API] Error procesando {blob.name}: {e}")
                    continue
        
        # Ordenar por fecha de procesamiento (más reciente primero)
        results.sort(key=lambda x: x['date'], reverse=True)

        # ── BigQuery fallback — carga datos históricos cuando GCS está vacío ──
        # Los registros legacy/missing_legacy no tienen archivos en GCS vault/;
        # solo existen en fact_kpi_values. Los sintetizamos en el mismo formato
        # financial_metrics_2025 que el frontend ya sabe consumir.
        if len(results) == 0:
            results = _build_results_from_bq(cid_clean)

        print(f" [API] Resultados encontrados para {company_id}: {len(results)}")

        return JSONResponse(
            content={
                "status": "success",
                "results": results,
                "company_id": company_id,
                "total": len(results)
            }
        )
    except (DefaultCredentialsError, Forbidden, Unauthorized) as e:
        sa_path = _resolve_service_account_path()
        print(f" [API] Error de credenciales/permisos obteniendo resultados: {e}")
        print(f"   GOOGLE_APPLICATION_CREDENTIALS={os.getenv('GOOGLE_APPLICATION_CREDENTIALS')}")
        print(f"   Service Account resuelto={sa_path}")
        raise HTTPException(
            status_code=500,
            detail="Error de autenticación/permisos con GCS. Verifica credenciales."
        )
    except Exception as e:
        print(f"[API] Error general en get_all_results: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo resultados: {str(e)}")

@app.get("/health")
async def health_check():
    """Endpoint de verificación de salud"""
    return {"status": "healthy", "service": "cometa-pipeline-api"}


# ── Analyst KPI correction ────────────────────────────────────────────────────

class KpiUpdateRequest(BaseModel):
    """
    Body for PUT /api/kpi-update.

    Fields
    ------
    submission_id : UUID of the submission in BigQuery (fact_kpi_values.submission_id).
    metric_id     : KPI key as stored in fact_kpi_values.kpi_key
                    (e.g. "revenue_growth", "ebitda_margin").
    value         : New raw string value typed by the Analyst
                    (e.g. "42%", "$8.5M"). Passed through parse_numeric.
    """
    submission_id: str
    metric_id:     str
    value:         str


@app.put("/api/kpi-update")
async def kpi_update(payload: KpiUpdateRequest, token: dict = Depends(_require_auth)):
    """
    Persist an Analyst correction to a KPI row in BigQuery.

    - Validates the new value via parse_numeric (Rule 4).
    - Sets is_manually_edited=TRUE and edited_at=CURRENT_TIMESTAMP().
    - Preserves the original value in edited_raw_value for audit.

    Returns the updated KPI data including parse result.
    """
    try:
        result = update_kpi_value(
            submission_id=payload.submission_id,
            kpi_key=payload.metric_id,
            new_raw_value=payload.value,
        )
        return JSONResponse(
            content={
                "status":  "success",
                "message": f"KPI '{payload.metric_id}' actualizado correctamente",
                **result,
            }
        )

    except ValueError as e:
        # Row not found in BigQuery
        raise HTTPException(status_code=404, detail=str(e))

    except Exception as e:
        print(f"❌ [API] Error en /api/kpi-update: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error actualizando KPI en BigQuery: {str(e)}"
        )


# ── PATCH /api/kpi/update — Analyst correction with fidelity upgrade ─────────
#
# Diferencias clave respecto al PUT /api/kpi-update:
#   1. Solo accesible por ANALISTA (ANA-* / role=ANALISTA) — 403 para FOUNDER.
#   2. Al ser editado manualmente, el view stg_legacy_fact_kpis retorna
#      value_status='verified' para ese row (is_manually_edited=TRUE es la
#      señal que usa el view). Esto hace que el Banner de Fidelidad desaparezca
#      conforme el analista limpia la data periodo a periodo.

@app.patch("/api/kpi/update")
async def kpi_patch_update(
    payload: KpiUpdateRequest,
    token: dict = Depends(_require_auth),
):
    """
    Corrige el valor de un KPI en BigQuery y lo marca como 'verified'.

    - Requiere rol ANALISTA (403 si es FOUNDER o SOCIO).
    - Llama a update_kpi_value() que establece is_manually_edited=TRUE,
      lo cual hace que stg_legacy_fact_kpis retorne value_status='verified'
      para ese row — haciendo desaparecer el Banner de Fidelidad cuando todos
      los KPIs del periodo estén verificados.

    Body
    ----
    submission_id : str  — UUID de la submision en BigQuery.
    metric_id     : str  — kpi_key (e.g. "revenue_growth").
    value         : str  — nuevo valor corregido (e.g. "42%", "$8.5M").
    """
    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_analista = (role == "ANALISTA") or user_id.startswith("ANA-")
    if not is_analista:
        raise HTTPException(
            status_code=403,
            detail="Solo los analistas pueden corregir KPIs.",
        )

    try:
        result = update_kpi_value(
            submission_id=payload.submission_id,
            kpi_key=payload.metric_id,
            new_raw_value=payload.value,
        )
        return JSONResponse(
            content={
                "status":       "success",
                "message":      f"KPI '{payload.metric_id}' corregido — fidelidad actualizada",
                "value_status": "verified",
                **result,
            }
        )

    except ValueError as e:
        raise HTTPException(status_code=404, detail=str(e))

    except Exception as e:
        print(f"❌ [API] Error en PATCH /api/kpi/update: {e}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error corrigiendo KPI en BigQuery: {str(e)}",
        )


# ── Analyst batch-edit with audit hash ───────────────────────────────────────

class AnalystEditRequest(BaseModel):
    """Body for POST /api/analyst/audit-edit."""
    submission_id: str                   # file_hash / submission_id in BigQuery
    updates:       dict[str, str]        # { kpi_key: new_raw_value }
    note:          str = ""              # Optional edit justification note


@app.post("/api/analyst/audit-edit")
@limiter.limit("30/minute")
async def analyst_audit_edit(
    request: Request,
    body: AnalystEditRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Batch-edit KPI values for a submission and return an audit hash.

    Only accessible to ANALISTA role.  For each entry in ``body.updates``,
    calls ``update_kpi_value()`` (sets ``is_manually_edited=TRUE`` + audit
    trail in ``edited_raw_value``).  After all updates a SHA-256 vault seal
    is generated covering: submission_id + sorted kpi_keys + analyst email +
    timestamp.  This hash is the "recibo de edición" the analyst can attach
    to the case.

    Returns
    -------
    JSON ``{ status, audit_hash, updated_kpis, failed_kpis, submission_id }``
    """
    from src.services.hash_service import generate_vault_seal

    # ── A1: ANALISTA-only gate ────────────────────────────────────────────────
    role    = token.get("role", "")
    user_id = token.get("user_id", "")
    if role != "ANALISTA" and not user_id.startswith("ANA-"):
        raise HTTPException(status_code=403, detail="Solo analistas pueden editar en batch.")

    analyst_email: str = (token.get("email") or token.get("sub", "")).strip()

    if not body.updates:
        raise HTTPException(status_code=422, detail="updates no puede estar vacío.")

    processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
    updated_kpis: list[str] = []
    failed_kpis:  list[dict] = []

    for kpi_key, raw_value in body.updates.items():
        raw_value = (raw_value or "").strip()
        if not raw_value:
            continue
        try:
            update_kpi_value(
                submission_id=body.submission_id,
                kpi_key=kpi_key,
                new_raw_value=raw_value,
            )
            updated_kpis.append(kpi_key)
        except Exception as _err:
            failed_kpis.append({"kpi_key": kpi_key, "error": str(_err)})
            print(f"⚠️  [analyst/audit-edit] {kpi_key}: {_err}")

    if not updated_kpis:
        raise HTTPException(
            status_code=422,
            detail=f"No se pudo actualizar ningún KPI. Detalles: {failed_kpis}",
        )

    # ── Generate audit hash ───────────────────────────────────────────────────
    audit_hash = generate_vault_seal(
        company_id   = analyst_email,
        file_hash    = body.submission_id,
        kpi_rows     = [
            {"kpi_key": k, "raw_value": body.updates[k], "unit": "", "is_valid": True}
            for k in sorted(updated_kpis)
        ],
        processed_at = processed_at,
    )

    print(
        f"[analyst/audit-edit] analyst={analyst_email!r}  "
        f"submission={body.submission_id[:12]}…  "
        f"updated={len(updated_kpis)}  failed={len(failed_kpis)}  "
        f"hash={audit_hash[:16]}…"
    )

    return JSONResponse(content={
        "status":       "ok",
        "audit_hash":   audit_hash,
        "updated_kpis": updated_kpis,
        "failed_kpis":  failed_kpis,
        "submission_id": body.submission_id,
        "processed_at": processed_at,
    })


# ── Coverage heatmap — ANALISTA only ─────────────────────────────────────────

@app.get("/api/analyst/coverage")
async def analyst_coverage(
    token:        dict = Depends(_require_auth),
    portfolio_id: str  = Query(default=""),
):
    """
    GET /api/analyst/coverage — Portfolio KPI coverage matrix.

    Returns per-company × per-period KPI verification status for the heatmap
    component.  Restricted to ANALISTA role.

    Response shape
    --------------
    {
        "status":    "ok",
        "companies": [{"key": str, "display": str, "portfolio_id": str}],
        "periods":   [str],      # canonical PYYYYQxMyy, chronological
        "cells":     [
            {
                "company":        str,
                "period":         str,
                "status":         "verified" | "legacy" | "missing",
                "kpi_count":      int,
                "verified_count": int,
                "legacy_count":   int
            }
        ]
    }
    """
    role    = (token.get("role") or "").upper()
    user_id = (token.get("user_id") or "")

    is_analista = (role == "ANALISTA") or user_id.startswith("ANA-")
    if not is_analista:
        raise HTTPException(
            status_code=403,
            detail="Acceso restringido a analistas Cometa.",
        )

    try:
        data = query_coverage(portfolio_id=portfolio_id.strip() or None)
        return JSONResponse({"status": "ok", **data})
    except Exception as exc:
        print(f"❌ [API] Error en GET /api/analyst/coverage: {exc}")
        traceback.print_exc()
        raise HTTPException(
            status_code=500,
            detail=f"Error consultando cobertura: {str(exc)}",
        )


# ── Manual KPI entry (Analista Auditoría tab) ────────────────────────────────

_KPI_ENTRY_FIELDS = frozenset({
    "revenue_growth", "gross_profit_margin", "ebitda_margin",
    "cash_in_bank_end_of_year", "annual_cash_flow", "working_capital_debt",
    "revenue", "ebitda", "cogs",
    "mrr", "churn_rate", "cac", "portfolio_size", "npl_ratio", "gmv", "loss_ratio",
})


class ManualEntryRequest(BaseModel):
    """Body for POST /api/manual-entry — Analyst enters KPIs without a PDF.

    Supports all 16 KPIs of the current data contract:
      - 6 core financial KPIs
      - 3 base metrics (inputs for the derivation engine)
      - 7 sector metrics (sector-specific KPIs)

    Empty strings from the frontend are coerced to None before validation.
    Numeric values sent as float/int are coerced to str for parse_numeric.
    """
    # ── Identifiers ────────────────────────────────────────────────────────
    company_id:               str
    portfolio_id:             str
    period_id:                str  = "FY2025"
    founder_email:            str  = ""
    submission_id:            str | None = None   # links manual entry to original upload
    # ── Core financial KPIs ────────────────────────────────────────────────
    revenue_growth:           str | None = None
    gross_profit_margin:      str | None = None
    ebitda_margin:            str | None = None
    cash_in_bank_end_of_year: str | None = None
    annual_cash_flow:         str | None = None
    working_capital_debt:     str | None = None
    # ── Base metrics (derivation engine inputs) ────────────────────────────
    revenue:                  str | None = None
    ebitda:                   str | None = None
    cogs:                     str | None = None
    # ── Sector metrics ─────────────────────────────────────────────────────
    mrr:                      str | None = None   # SAAS
    churn_rate:               str | None = None   # SAAS
    cac:                      str | None = None   # SAAS / ECOM / INSUR
    portfolio_size:           str | None = None   # LEND
    npl_ratio:                str | None = None   # LEND
    gmv:                      str | None = None   # ECOM
    loss_ratio:               str | None = None   # INSUR

    @model_validator(mode="before")
    @classmethod
    def _coerce_kpi_fields(cls, data: dict) -> dict:
        """Convert empty strings → None and numbers → str for all KPI fields."""
        if not isinstance(data, dict):
            return data
        for field in _KPI_ENTRY_FIELDS:
            val = data.get(field)
            if val is None:
                continue
            if isinstance(val, (int, float)):
                # Frontend sent a bare number — convert to string for parse_numeric
                data[field] = str(val)
            elif isinstance(val, str) and val.strip() == "":
                # Empty string — treat as not provided
                data[field] = None
        return data


def _manual_node(value: str | None) -> dict:
    return {
        "value":       value,
        "confidence":  1.0,
        "description": "Entrada manual del Analista",
    }


@app.post("/api/manual-entry")
async def manual_entry(payload: ManualEntryRequest, token: dict = Depends(_require_auth)):
    """
    Persist analyst-entered KPIs directly to BigQuery without a PDF.
    Builds a synthetic Gemini JSON, runs it through build_contract(),
    and inserts it via insert_contract() — same deduplication rules apply.
    """
    import hashlib, uuid as _uuid

    # Synthetic Gemini JSON mirrors the real 16-KPI schema exactly
    synthetic_gemini = {
        "_document_context": {
            "currency":    "USD",
            "period":      payload.period_id,
            "scale":       "units",
            "scale_notes": "Entrada manual del Analista",
        },
        "financial_metrics_2025": {
            # ── Core KPIs ──────────────────────────────────────────────────
            "revenue_growth": _manual_node(payload.revenue_growth),
            "profit_margins": {
                "gross_profit_margin": _manual_node(payload.gross_profit_margin),
                "ebitda_margin":       _manual_node(payload.ebitda_margin),
            },
            "cash_flow_indicators": {
                "cash_in_bank_end_of_year": _manual_node(payload.cash_in_bank_end_of_year),
                "annual_cash_flow":         _manual_node(payload.annual_cash_flow),
            },
            "debt_ratios": {
                "working_capital_debt": _manual_node(payload.working_capital_debt),
            },
            # ── Base metrics (OBS-03) ──────────────────────────────────────
            "base_metrics": {
                "revenue": _manual_node(payload.revenue),
                "ebitda":  _manual_node(payload.ebitda),
                "cogs":    _manual_node(payload.cogs),
            },
            # ── Sector metrics (OBS-03) ────────────────────────────────────
            "sector_metrics": {
                "mrr":           _manual_node(payload.mrr),
                "churn_rate":    _manual_node(payload.churn_rate),
                "cac":           _manual_node(payload.cac),
                "portfolio_size":_manual_node(payload.portfolio_size),
                "npl_ratio":     _manual_node(payload.npl_ratio),
                "gmv":           _manual_node(payload.gmv),
                "loss_ratio":    _manual_node(payload.loss_ratio),
            },
        },
    }

    # Unique hash based on company + period + timestamp so re-entries don't dedup
    _company_id    = payload.company_id
    _period_id     = payload.period_id
    _unique_salt   = str(_uuid.uuid4())
    raw_hash_input = f"{_company_id}:{_period_id}:{_unique_salt}"
    file_hash = hashlib.sha256(raw_hash_input.encode()).hexdigest()[:16]

    try:
        contract = build_contract(
            gemini_json=synthetic_gemini,
            file_hash=file_hash,
            company_id=payload.company_id,
            founder_email=payload.founder_email,
            original_filename=f"[manual] {payload.company_id} {payload.period_id}",
            portfolio_id=payload.portfolio_id,
        )
        db_result = insert_contract(contract)
        return JSONResponse(content={
            "status":       "success",
            "message":      f"Datos de {payload.company_id} guardados correctamente",
            "submission":   contract["submission"],
            "kpi_rows":     contract["kpi_rows"],
            "db":           db_result,
        })
    except Exception as e:
        print(f"[API] Error en /api/manual-entry: {e}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))


# ── Duplicate audit — delete a vault submission ───────────────────────────────

@app.delete("/api/submission")
async def delete_submission(file_hash: str, company_id: str, token: dict = Depends(_require_auth)):
    """
    Delete a specific submission from the GCS vault.
    Identifies the blob by matching file_hash in metadata.
    BigQuery row is NOT deleted (preserves audit trail).
    """
    try:
        storage_client = _get_storage_client()
        bucket         = storage_client.bucket(GCS_OUTPUT_BUCKET)
        vault_prefix   = f"vault/{company_id}/"
        blobs          = list(bucket.list_blobs(prefix=vault_prefix))
        deleted        = 0
        for blob in blobs:
            if blob.metadata and blob.metadata.get("file_hash") == file_hash:
                blob.delete()
                deleted += 1
                print(f"[API] Deleted vault blob: {blob.name}")
        if deleted == 0:
            raise HTTPException(status_code=404, detail=f"No blob found for hash {file_hash} in {company_id}")
        return JSONResponse(content={"status": "success", "deleted": deleted, "file_hash": file_hash})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


# ── Portfolio registry endpoints ──────────────────────────────────────────────

@app.get("/api/kpi-metadata")
async def get_kpi_metadata(vertical: str | None = None):
    """
    Returns the KPI dictionary from dim_kpi_metadata.

    Query params
    ------------
    vertical : optional — 'SAAS' | 'FINTECH' | 'MARKETPLACE' | 'INSURTECH'
               When provided, returns GENERAL KPIs plus vertical-specific ones.
               When omitted, returns the full catalogue.

    No auth required: founders need this before they are logged in to
    the analyst cockpit (UploadFlow step 0 vertical selector).
    """
    try:
        items = query_kpi_metadata(vertical)
    except Exception as e:
        print(f"⚠️  [kpi-metadata] Non-fatal BQ error, using seed fallback: {e}")
        items = []
    return JSONResponse(content={"status": "ok", "kpis": items, "vertical": vertical})


@app.get("/api/portfolio-companies")
async def get_portfolio_companies(portfolio_id: str = None):
    """
    Returns the canonical company list grouped by portfolio.
    Each company includes has_data: bool — True when fact_kpi_values has at
    least one row with value_status in ('legacy', 'missing_legacy', 'verified').

    Query params
    ------------
    portfolio_id : optional — "VII" or "CIII". If omitted, returns all funds.
    """
    # ── Fetch companies that have data in BQ (non-fatal) ──────────────────────
    companies_with_data: set[str] = set()
    try:
        _bq   = _get_bq_client_for_api()
        _ds   = os.getenv("BIGQUERY_DATASET", "cometa_vault_test")
        _sql  = f"""
            SELECT DISTINCT company_id
            FROM `{_bq.project}.cometa_vault.stg_legacy_fact_kpis`
            WHERE value_status IN ('legacy', 'missing_legacy', 'verified')
        """
        for row in _bq.query(_sql).result():
            if row.company_id:
                companies_with_data.add(row.company_id.lower())
    except Exception as _bq_err:
        # BQ unavailable — default has_data=True so UI stays usable
        print(f"[portfolio-companies] BQ check failed (non-fatal): {_bq_err}")
        companies_with_data = None  # type: ignore[assignment]

    # grouped: portfolio_id → list of {key, label, is_overview, has_data}
    grouped: dict[str, list[dict]] = {}
    for key, info in PORTFOLIO_MAP.items():
        pid = info["portfolio_id"]
        if portfolio_id and pid != portfolio_id:
            continue
        label       = info.get("display_name") or key.capitalize()
        is_overview = info.get("entity_type") == "FUND_OVERVIEW"
        has_data    = True if companies_with_data is None else (key.lower() in companies_with_data)
        grouped.setdefault(pid, []).append({
            "key":         key,
            "label":       label,
            "is_overview": is_overview,
            "has_data":    has_data,
        })

    def _sort_entries(entries: list[dict]) -> list[dict]:
        # Fund overviews float to the top, rest sorted alpha by label
        overviews = [e for e in entries if e["is_overview"]]
        rest      = sorted((e for e in entries if not e["is_overview"]), key=lambda e: e["label"])
        return overviews + rest

    return JSONResponse(content={
        "status":    "success",
        "portfolios": [
            {
                "portfolio_id":   pid,
                "portfolio_name": f"Fondo {pid}",
                "companies":      _sort_entries(entries),
            }
            for pid, entries in sorted(grouped.items())
        ],
    })


@app.get("/api/results/all")
async def get_all_results_global(token: dict = Depends(_require_auth)):
    """
    Returns the latest result per company across the whole portfolio.
    Source priority: GCS vault (uploads) → BigQuery legacy fallback.
    Companies only in BQ are returned with _source='bigquery_legacy'.
    """
    try:
        storage_client = _get_storage_client()
        bucket         = storage_client.bucket(GCS_OUTPUT_BUCKET)

        # ── 1. GCS scan — most-recent result per company ──────────────────────
        gcs_by_company: dict[str, dict] = {}   # company_id → result_item
        for blob in bucket.list_blobs(prefix="vault/"):
            if not blob.name.endswith(".json"):
                continue
            try:
                content = blob.download_as_text()
                try:
                    result_data = json.loads(content)
                    if isinstance(result_data, str):
                        result_data = json.loads(result_data)
                except json.JSONDecodeError:
                    continue
                if not isinstance(result_data, dict):
                    continue

                metadata   = blob.metadata or {}
                parts      = blob.name.split("/")
                company_id = parts[1] if len(parts) >= 3 else "unknown"
                processed  = metadata.get("processed_at", "unknown")

                existing = gcs_by_company.get(company_id)
                if existing and existing["date"] >= processed:
                    continue  # keep newer

                vault_prefix = f"vault/{company_id}/"
                gcs_by_company[company_id] = {
                    "id":   blob.name.replace(".json", "").replace(vault_prefix, ""),
                    "date": processed,
                    "data": result_data,
                    "metadata": {
                        "original_filename": metadata.get("original_filename", "unknown"),
                        "founder_email":     metadata.get("founder_email",     "unknown"),
                        "file_hash":         metadata.get("file_hash",         ""),
                        "processed_at":      processed,
                        "gcs_path":          blob.name,
                        "portfolio_id":      metadata.get("portfolio_id") or lookup_portfolio(
                                                metadata.get("company_domain", company_id)),
                        "company_domain":    company_id,
                    },
                }
            except Exception as exc:
                print(f"[API/all] GCS error on {blob.name}: {exc}")

        print(f"[API/all] GCS: {len(gcs_by_company)} empresas")

        # ── 2. BigQuery fallback — latest period for companies not in GCS ─────
        bq_results = _build_all_results_from_bq(set(gcs_by_company.keys()))

        # ── 3. Merge: GCS first, then BQ for the rest ─────────────────────────
        results = list(gcs_by_company.values()) + bq_results
        results.sort(key=lambda x: x["date"], reverse=True)

        print(f"[API/all] Total: {len(results)} ({len(gcs_by_company)} GCS + {len(bq_results)} BQ legacy)")

        return JSONResponse(content={
            "status":  "success",
            "results": results,
            "total":   len(results),
        })
    except Exception as e:
        print(f"[API/all] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo todos los resultados: {str(e)}")


@app.get("/api/analytics/portfolio")
async def get_portfolio_analytics(portfolio_id: str = "CIII", token: dict = Depends(_require_auth)):
    """
    Aggregated KPI analytics from BigQuery for the requested portfolio.
    Groups by (month, company_id) and returns per-KPI averages.
    """
    try:
        result = query_portfolio_analytics(portfolio_id)
        return JSONResponse(content={"status": "success", **result})
    except Exception as e:
        print(f"[API/analytics] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error obteniendo analytics: {str(e)}")


@app.get("/api/audit")
async def get_audit_report(portfolio_id: str = None):
    """
    Runs the BigQuery post-insert audit across fact_kpi_values.
    Returns all rows flagged as ERROR or WARNING with their audit_status.

    Query params
    ------------
    portfolio_id : optional — "VII" or "CIII". If omitted, audits all funds.

    audit_status values
    -------------------
    PASS                          — row is clean
    ERROR: Duplicado              — more than one row for same (company, metric, period)
    ERROR: Valor no numérico      — is_valid = FALSE
    ERROR: Confianza crítica (<0.70) — Gemini was highly uncertain
    ADVERTENCIA: Confianza baja (<0.85) — flagged for human review
    """
    try:
        result = run_audit_query(portfolio_id=portfolio_id)
        return JSONResponse(content={"status": "success", **result})
    except Exception as e:
        print(f"[API/audit] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error ejecutando audit: {str(e)}")


@app.get("/api/audit/fidelity/{submission_id}")
async def get_fidelity_audit(submission_id: str):
    """
    Reporte de Fidelidad de Datos — Auditor Senior Cometa.

    Ejecuta tres auditorías encadenadas sobre una submission específica:

    1. identity_check
       Verifica que company_id esté en los 30 registros oficiales de dim_company
       y que el bucket_id asignado coincida con el registro canónico COMPANY_BUCKET.

    2. calculator_audit
       Clasifica cada KPI como 'gemini' (extraído del PDF) o 'calculated' (derivado
       por Python).  Para gross_profit_margin y ebitda_margin re-ejecuta la fórmula
       matemática y compara contra el valor almacenado.  Discrepancias > 0.5pp
       generan WARN; > 2pp generan ERROR (posible manipulación de reporte founder).

    3. checklist_diagnosis
       Cruza los KPIs válidos del reporte con SECTOR_REQUIREMENTS del vertical
       (SAAS / LEND / ECOM / INSUR / OTH) y devuelve missing_kpis con mensaje
       listo para mostrar al founder.

    overall_status
    --------------
    PASS  — sin hallazgos de error ni advertencia
    WARN  — advertencias presentes; el reporte es usable pero debe revisarse
    FAIL  — errores bloqueantes detectados (identidad incorrecta, discrepancia alta)

    Path param
    ----------
    submission_id : UUID de la submission (devuelto por /upload o /api/submissions)
    """
    try:
        result = run_fidelity_audit(submission_id)
        return JSONResponse(content={"status": "success", **result}, status_code=200)
    except ValueError as e:
        # OBS-02: submission_id no existe en BigQuery → 404 real
        raise HTTPException(status_code=404, detail=str(e))
    except Exception as e:
        print(f"[API/fidelity-audit] Error: {e}")
        raise HTTPException(status_code=500, detail=f"Error ejecutando fidelity audit: {str(e)}")


# ── CSV Export endpoint ───────────────────────────────────────────────────────

import csv as _csv_mod
import io  as _io_mod

from fastapi.responses import Response as _PlainResponse

@app.get("/api/export/csv")
@limiter.limit("10/minute")
async def export_kpi_csv(
    request:      Request,
    portfolio_id: str | None = None,
    company_id:   str | None = None,
    token:        dict = Depends(_require_auth),
    _origin:      None = Depends(_verify_origin),
):
    """
    Export KPI data as UTF-8 CSV from BigQuery.

    Query params
    ------------
    portfolio_id : filter by fund  (e.g. "CIII" or "VII")
    company_id   : filter by company domain  (e.g. "solvento.com")

    Multi-tenant rule (A1):
    - FND- founders are hard-locked to their own company_id from JWT.
    - ANA- analysts may export any scope requested.

    Columns: Empresa, Fondo, Período, KPI, Valor, Unidad, Confianza, Procesado
    """
    from google.cloud import bigquery as _bq

    # A1 — tenant isolation
    jwt_company      = _derive_tenant_from_token(token)
    effective_company = jwt_company if jwt_company is not None else company_id

    # Build parameterised BigQuery query
    ds      = f"{PROJECT_ID}.{os.getenv('BIGQUERY_DATASET', 'cometa_vault')}"
    filters = ["f.is_valid = TRUE", "f.raw_value IS NOT NULL"]
    params: list[_bq.ScalarQueryParameter] = []

    if portfolio_id:
        filters.append("s.portfolio_id = @portfolio_id")
        params.append(_bq.ScalarQueryParameter("portfolio_id", "STRING", portfolio_id))
    if effective_company:
        filters.append("LOWER(s.company_id) LIKE @company_id")
        params.append(_bq.ScalarQueryParameter("company_id", "STRING",
                                               f"%{effective_company.lower()}%"))

    sql = f"""
        SELECT
            s.company_id,
            s.portfolio_id,
            s.period_id,
            f.kpi_label,
            f.raw_value,
            f.unit,
            f.confidence_score,
            s.submitted_at
        FROM `{ds}.fact_kpi_values` f
        JOIN `{ds}.submissions`     s ON f.submission_id = s.submission_id
        WHERE {' AND '.join(filters)}
        ORDER BY s.company_id, s.submitted_at DESC, f.kpi_label
        LIMIT 10000
    """

    try:
        client  = _get_bq_client_for_api()
        job     = client.query(sql, job_config=_bq.QueryJobConfig(query_parameters=params))
        rows    = list(job.result())
    except Exception as exc:
        print(f"[export/csv] BigQuery error: {exc}")
        raise HTTPException(status_code=500,
                            detail=f"Error consultando BigQuery: {str(exc)}")

    # Build CSV in memory
    buf    = _io_mod.StringIO()
    writer = _csv_mod.writer(buf)
    writer.writerow(["Empresa", "Fondo", "Período", "KPI", "Valor",
                     "Unidad", "Confianza", "Procesado"])
    for r in rows:
        conf_str = f"{float(r.confidence_score):.2f}" if r.confidence_score is not None else ""
        date_str = str(r.submitted_at)[:10] if r.submitted_at else ""
        writer.writerow([
            r.company_id   or "",
            r.portfolio_id or "",
            r.period_id    or "",
            r.kpi_label    or "",
            r.raw_value    or "",
            r.unit         or "",
            conf_str,
            date_str,
        ])

    scope_tag = (effective_company or portfolio_id or "all").replace("/", "-")
    filename  = f"cometa_kpis_{scope_tag}_{datetime.now().strftime('%Y%m%d')}.csv"

    return _PlainResponse(
        content=buf.getvalue().encode("utf-8-sig"),   # BOM for Excel compatibility
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── RAG Chat endpoint ─────────────────────────────────────────────────────────

from src.ai_engine import (
    build_rag_prompt,
    call_gemini        as _call_gemini_engine,
    call_gemini_stream as _call_gemini_stream,
)

_CHAT_MAX_QUESTION_CHARS   = 500
_CHAT_MAX_SUMMARY_CHARS    = 300  # cap the frontend-supplied executive summary

class ChatRequest(BaseModel):
    """
    Body for POST /api/chat.

    Fields
    ------
    question:           The analyst's question (max 500 chars — C5).
    portfolio_id:       Optional portfolio filter.
    company_id:         Company in focus — only respected for ANA- users (A1).
    executive_summary:  Pre-computed KPI one-liner from ExecutiveSummaryText;
                        injected into the Gemini prompt when caller is ANA-.
    """
    question:           str
    portfolio_id:       str | None = None
    company_id:         str | None = None       # ignored for FND- founders (A1)
    executive_summary:  str | None = None       # ANA- analyst context only

@app.post("/api/chat")
@limiter.limit("20/minute")
async def portfolio_chat(
    request: Request,
    req: ChatRequest,
    token: dict = Depends(_require_auth),
    _origin: None = Depends(_verify_origin),
):
    """
    RAG chat: consulta BigQuery → arma contexto → llama Gemini → devuelve respuesta.
    Controles activos: C5 (prompt injection), A1 (multi-tenant), A2 (rate limit), C4 (origin).
    """
    # C5 — Límite de longitud de la pregunta
    question_raw = (req.question or "").strip()
    if not question_raw:
        raise HTTPException(status_code=400, detail="La pregunta no puede estar vacía.")
    if len(question_raw) > _CHAT_MAX_QUESTION_CHARS:
        raise HTTPException(
            status_code=400,
            detail=f"La pregunta excede el límite de {_CHAT_MAX_QUESTION_CHARS} caracteres "
                   f"({len(question_raw)} recibidos).",
        )

    # A1 — Aislamiento multi-tenant: company_id derivado del JWT, nunca del body
    jwt_company_id       = _derive_tenant_from_token(token)
    effective_company_id = jwt_company_id if jwt_company_id is not None else req.company_id

    # I1 — Extracción de identidad del JWT (Capa Humana)
    user_id_claim: str = token.get("user_id", "")
    user_name:     str = (token.get("name")  or token.get("sub") or "").strip()
    user_role:     str = (token.get("role")  or "").strip()
    is_analyst         = user_id_claim.startswith("ANA-")

    # C5 — Cap the executive summary length to avoid prompt stuffing
    executive_summary: str | None = None
    if is_analyst and req.executive_summary:
        executive_summary = req.executive_summary.strip()[:_CHAT_MAX_SUMMARY_CHARS]

    # 1. Recuperar contexto de BigQuery
    raw_rows = _query_rag_context(req.portfolio_id, effective_company_id)

    # A3 — RAG Leak Protection: verifica que todos los rows pertenezcan al tenant
    rows = _verify_rag_integrity(raw_rows, effective_company_id or "")

    # Detectar si algún KPI en el contexto aún no ha sido verificado manualmente
    has_legacy_data = any(not row.get("is_manually_edited", False) for row in rows)

    # KPI Dictionary — enriquece el prompt con definiciones y año de alta (non-fatal)
    kpi_dict = _fetch_kpi_dict_for_rag()

    # 2. Build structured prompt via ai_engine — con identidad, advertencia legacy y diccionario
    prompt = build_rag_prompt(
        question=question_raw,
        context_rows=rows,
        company_id=effective_company_id,
        portfolio_id=req.portfolio_id,
        executive_summary=executive_summary,
        is_analyst=is_analyst,
        user_name=user_name,
        user_role=user_role,
        has_legacy_data=has_legacy_data,
        kpi_dict=kpi_dict,
    )

    # 3. Llamar Gemini via ai_engine
    try:
        answer = _call_gemini_engine(prompt, PROJECT_ID, VERTEX_LOCATION)
    except Exception as e:
        print(f"[RAG/chat] Gemini error: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando respuesta: {str(e)}")

    # 4. AI Audit Log — registrar la consulta (non-fatal)
    try:
        insert_ai_audit_log(
            user_id         = user_id_claim,
            user_name       = user_name,
            user_role       = user_role,
            question        = question_raw,
            context_rows    = len(rows),
            has_legacy_data = has_legacy_data,
            endpoint        = "/api/chat",
            company_id      = effective_company_id or "",
            portfolio_id    = req.portfolio_id or "",
        )
    except Exception as _audit_err:
        print(f"⚠️  [RAG/audit] Non-fatal audit log error: {_audit_err}")

    return JSONResponse(content={
        "status":        "success",
        "answer":        answer,
        "sources_count": len(rows),
        "has_legacy_data": has_legacy_data,
        "portfolio_id":  req.portfolio_id,
        "company_id":    effective_company_id,
    })


@app.post("/api/chat/stream")
@limiter.limit("20/minute")
async def portfolio_chat_stream(
    request: Request,
    req: ChatRequest,
    token: dict = Depends(_require_auth),
    _origin: None = Depends(_verify_origin),
):
    """
    SSE streaming chat — identical security controls as /api/chat.

    Response format (text/event-stream):
      data: {"token": "<chunk>"}\\n\\n   — incremental Gemini output
      data: {"error":  "<msg>"}\\n\\n   — on Gemini error
      data: [DONE]\\n\\n               — stream completed
    """
    import json as _json_mod

    # ── Input validation (mirrors /api/chat) ─────────────────────────────────
    question_raw = (req.question or "").strip()
    if not question_raw:
        raise HTTPException(status_code=400, detail="La pregunta no puede estar vacía.")
    if len(question_raw) > _CHAT_MAX_QUESTION_CHARS:
        raise HTTPException(
            status_code=400,
            detail=f"La pregunta excede {_CHAT_MAX_QUESTION_CHARS} caracteres.",
        )

    # ── A1 — tenant isolation ──────────────────────────────────────────────────
    jwt_company_id       = _derive_tenant_from_token(token)
    effective_company_id = jwt_company_id if jwt_company_id is not None else req.company_id

    # I1 — Extracción de identidad del JWT (Capa Humana)
    user_id_claim: str = token.get("user_id", "")
    user_name:     str = (token.get("name")  or token.get("sub") or "").strip()
    user_role:     str = (token.get("role")  or "").strip()
    is_analyst         = user_id_claim.startswith("ANA-")

    executive_summary: str | None = None
    if is_analyst and req.executive_summary:
        executive_summary = req.executive_summary.strip()[:_CHAT_MAX_SUMMARY_CHARS]

    # ── Build prompt with A3 leak protection ──────────────────────────────────
    raw_rows = _query_rag_context(req.portfolio_id, effective_company_id)
    rows     = _verify_rag_integrity(raw_rows, effective_company_id or "")

    has_legacy_data = any(not row.get("is_manually_edited", False) for row in rows)

    kpi_dict = _fetch_kpi_dict_for_rag()

    prompt = build_rag_prompt(
        question=question_raw,
        context_rows=rows,
        company_id=effective_company_id,
        portfolio_id=req.portfolio_id,
        executive_summary=executive_summary,
        is_analyst=is_analyst,
        user_name=user_name,
        user_role=user_role,
        has_legacy_data=has_legacy_data,
        kpi_dict=kpi_dict,
    )

    # ── AI Audit Log — registrar antes de iniciar el stream (non-fatal) ────────
    try:
        insert_ai_audit_log(
            user_id         = user_id_claim,
            user_name       = user_name,
            user_role       = user_role,
            question        = question_raw,
            context_rows    = len(rows),
            has_legacy_data = has_legacy_data,
            endpoint        = "/api/chat/stream",
            company_id      = effective_company_id or "",
            portfolio_id    = req.portfolio_id or "",
        )
    except Exception as _audit_err:
        print(f"⚠️  [RAG/audit/stream] Non-fatal audit log error: {_audit_err}")

    # ── SSE generator ──────────────────────────────────────────────────────────
    async def _sse_generator():
        try:
            for chunk in _call_gemini_stream(prompt, PROJECT_ID, VERTEX_LOCATION):
                payload = _json_mod.dumps({"token": chunk}, ensure_ascii=False)
                yield f"data: {payload}\n\n"
        except Exception as exc:
            err = _json_mod.dumps({"error": str(exc)}, ensure_ascii=False)
            yield f"data: {err}\n\n"
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        _sse_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control":     "no-cache",
            "X-Accel-Buffering": "no",   # disables nginx proxy buffering
        },
    )


# ═══════════════════════════════════════════════════════════════════════════════
# AUTENTICACIÓN — /api/login  &  /api/me
# ═══════════════════════════════════════════════════════════════════════════════

class LoginRequest(BaseModel):
    email:    str
    password: str


# ── Helpers de contraseña (bcrypt) ────────────────────────────────────────────

def _is_bcrypt_hash(value: str) -> bool:
    """Detecta si `value` es un hash bcrypt válido ($2b$, $2a$ o $2y$)."""
    return isinstance(value, str) and value.startswith(("$2b$", "$2a$", "$2y$"))


def _hash_password(plaintext: str) -> str:
    """Genera un hash bcrypt con salt aleatorio (12 rondas por defecto)."""
    return bcrypt.hashpw(plaintext.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")


def _verify_password(plaintext: str, stored: str) -> bool:
    """
    Verifica la contraseña contra el valor almacenado.
    - Hash bcrypt → bcrypt.checkpw (timing-safe)
    - Texto plano legacy → comparación directa (solo durante migración)
    """
    if _is_bcrypt_hash(stored):
        return bcrypt.checkpw(plaintext.encode("utf-8"), stored.encode("utf-8"))
    return stored == plaintext


def _load_users() -> list[dict]:
    """Lee users.json desde disco. No cachea para reflejar cambios en caliente."""
    try:
        with open(_USERS_FILE, "r", encoding="utf-8") as fh:
            return json.load(fh).get("users", [])
    except FileNotFoundError:
        return []


def _save_users(users: list[UserSchema]) -> None:
    """
    Persiste usuarios validados en users.json (escritura atómica vía .tmp).

    La firma `list[UserSchema]` es la barrera de seguridad principal:
    es imposible llamar esta función con datos sin validar, ya que
    UserSchema aplica todas sus validaciones en el momento de construcción.

    Flujo garantizado:
      1. Caller construye list[UserSchema]  ← validación ocurre AQUÍ
      2. Si falla → ValidationError antes de abrir ningún archivo
      3. Si pasa → serialización + write atómico (.tmp → replace)
    """
    tmp = _USERS_FILE.with_suffix(".json.tmp")
    payload = [u.model_dump() for u in users]
    tmp.write_text(json.dumps({"users": payload}, indent=2, ensure_ascii=False), encoding="utf-8")
    tmp.replace(_USERS_FILE)


@app.post("/api/login")
@limiter.limit("10/minute")
async def login(request: Request, body: LoginRequest):
    """
    Valida credenciales contra users.json y emite un JWT de 24 h.

    Flujo de seguridad:
      1. Buscar usuario por email.
      2. Verificar contraseña con bcrypt.checkpw (o comparación directa si legacy).
      3. Migración perezosa si cualquier campo está desactualizado:
         - Contraseña en texto plano → hashear con bcrypt antes de guardar.
         - ID en formato legacy      → generar ID Híbrido.
         - Todo el archivo validado con UserSchema ANTES de la escritura atómica.
      4. Emitir JWT con user_id para auditoría.
    """
    email_lc = (body.email or "").strip().lower()
    users     = _load_users()

    idx  = next((i for i, u in enumerate(users) if u.get("email", "").lower() == email_lc), None)
    user = users[idx] if idx is not None else None

    # ── Verificación de credenciales ──────────────────────────────────────────
    if not user or not _verify_password(body.password, user.get("password", "")):
        raise HTTPException(status_code=401, detail="Credenciales inválidas")

    # ── Bloquear cuentas pendientes de activación ──────────────────────────────
    if user.get("status") == "PENDING_INVITE":
        raise HTTPException(
            status_code=401,
            detail="Cuenta pendiente de activación. Revisa tu correo para configurar tu acceso.",
        )

    # ── Migración perezosa: ID Híbrido + hash bcrypt ──────────────────────────
    needs_id_migration = not is_hybrid_id(user.get("id", ""))
    needs_pw_migration = not _is_bcrypt_hash(user.get("password", ""))

    if needs_id_migration or needs_pw_migration:
        if needs_id_migration:
            users[idx]["id"] = generate_hybrid_id(email_lc)
        if needs_pw_migration:
            # Hashear antes de validar con UserSchema (garantía: campo no vacío)
            users[idx]["password"] = _hash_password(body.password)

        # Construir list[UserSchema] — validación completa ANTES de abrir el disco.
        # Migración perezosa de cualquier otro usuario legacy en el mismo archivo.
        # Si falla → ValidationError → handler global 422, disco intacto.
        validated: list[UserSchema] = [
            UserSchema.model_validate(
                u if (is_hybrid_id(u.get("id", "")) and _is_bcrypt_hash(u.get("password", "")))
                else {
                    **u,
                    "id":       u["id"] if is_hybrid_id(u.get("id", ""))
                                else generate_hybrid_id(u.get("email", "")),
                    "password": u["password"] if _is_bcrypt_hash(u.get("password", ""))
                                else _hash_password(u.get("password", "")),
                }
            )
            for u in users
        ]

        _save_users(validated)    # escritura atómica: .tmp → replace
        user = users[idx]         # refrescar referencia local

    user_id = user["id"]
    role    = enforce_internal_role(email_lc, user.get("role", "FOUNDER"))
    name    = user.get("name", "")

    token = create_access_token(
        email=email_lc,
        role=role,
        name=name,
        user_id=user_id,
    )

    return {
        "access_token": token,
        "token_type":   "bearer",
        "user": {
            "user_id":    user_id,
            "email":      email_lc,
            "name":       name,
            "role":       role,
            "company_id": user.get("company_id", ""),
        },
    }


@app.get("/api/me")
async def get_me(token: dict = Depends(_require_auth)):
    """
    Endpoint protegido. Decodifica el JWT del header Authorization
    y devuelve los datos del usuario autenticado, incluyendo user_id para auditoría.
    """
    return {
        "user_id":    token.get("user_id", ""),
        "email":      token.get("email") or token.get("sub", ""),
        "name":       token.get("name", ""),
        "role":       token.get("role", ""),
        "company_id": token.get("company_id", ""),
    }


# ── Invite / Setup-password flow ─────────────────────────────────────────────

_INVITE_TOKEN_TYPE = "invite"
_INVITE_EXPIRE_HOURS = 48

# Regex: min 8 chars, at least one digit, at least one non-alphanumeric char
_PASSWORD_RE = re.compile(r"^(?=.*\d)(?=.*[\W_]).{8,}$")


class SetupPasswordRequest(BaseModel):
    """Body for POST /api/auth/setup-password."""
    token:            str
    password:         str
    password_confirm: str


@app.post("/api/auth/setup-password")
@limiter.limit("10/minute")
async def setup_password(
    request: Request,
    body: SetupPasswordRequest,
) -> JSONResponse:
    """
    Activate a PENDING_INVITE account by setting the initial password.

    Flow
    ----
    1. Decode & verify the invite JWT (type="invite", not expired).
    2. Find the user in users.json — must have status=PENDING_INVITE.
    3. Validate password strength (≥8 chars, ≥1 digit, ≥1 symbol).
    4. Hash password with bcrypt, set status="ACTIVE".
    5. Atomic write via _save_users().
    6. Issue a 24-h access JWT so the founder is logged in immediately.

    Returns
    -------
    Same shape as POST /api/login: { access_token, token_type, user }.
    """
    # ── 1. Validate invite token ───────────────────────────────────────────────
    try:
        claims = jwt.decode(
            body.token,
            _JWT_SECRET,
            algorithms=[_JWT_ALGORITHM],
            options={"verify_aud": False},
        )
    except JWTError as exc:
        raise HTTPException(status_code=400, detail=f"Token inválido o expirado: {exc}")

    if claims.get("type") != _INVITE_TOKEN_TYPE:
        raise HTTPException(status_code=400, detail="Token no es de tipo invitación")

    invite_email: str = (claims.get("sub") or "").strip().lower()
    if not invite_email:
        raise HTTPException(status_code=400, detail="Token no contiene email")

    # ── 2. Find pending user ───────────────────────────────────────────────────
    users = _load_users()
    idx   = next(
        (i for i, u in enumerate(users) if u.get("email", "").lower() == invite_email),
        None,
    )
    if idx is None:
        raise HTTPException(status_code=404, detail="Usuario no encontrado")

    user = users[idx]
    if user.get("status") != "PENDING_INVITE":
        raise HTTPException(
            status_code=409,
            detail="Esta cuenta ya está activa. Inicia sesión normalmente.",
        )

    # ── 3. Validate password strength ─────────────────────────────────────────
    if body.password != body.password_confirm:
        raise HTTPException(status_code=422, detail="Las contraseñas no coinciden")

    if not _PASSWORD_RE.match(body.password):
        raise HTTPException(
            status_code=422,
            detail="La contraseña debe tener al menos 8 caracteres, un número y un símbolo.",
        )

    # ── 4 & 5. Hash + atomic save ─────────────────────────────────────────────
    users[idx]["password"] = _hash_password(body.password)
    users[idx]["status"]   = "ACTIVE"

    validated: list[UserSchema] = [UserSchema.model_validate(u) for u in users]
    _save_users(validated)

    # ── 6. Issue access token ──────────────────────────────────────────────────
    activated = users[idx]
    role      = enforce_internal_role(invite_email, activated.get("role", "FOUNDER"))
    token     = create_access_token(
        email=invite_email,
        role=role,
        name=activated.get("name", ""),
        user_id=activated.get("id", ""),
    )
    print(f"[setup-password] Account activated: {invite_email}")

    return JSONResponse(
        content={
            "access_token": token,
            "token_type":   "bearer",
            "user": {
                "user_id":    activated.get("id", ""),
                "email":      invite_email,
                "name":       activated.get("name", ""),
                "role":       role,
                "company_id": activated.get("company_id", ""),
            },
        },
        status_code=200,
    )


# ── Founder notification endpoints ───────────────────────────────────────────

class NotifyUploadRequest(BaseModel):
    """Body for POST /api/notify/upload."""
    founder_email:  str
    file_hash:      str
    company_domain: str = ""


@app.post("/api/notify/upload")
@limiter.limit("20/minute")
async def notify_upload(
    request: Request,
    body: NotifyUploadRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Best-effort upload notification hook.

    Called fire-and-forget by the frontend after each successful document
    upload.  Logs the event; in production this is where a real-time
    Slack/Teams notification could be triggered.

    Always returns 200 so transient failures never block the UI.
    """
    email = (token.get("email") or token.get("sub", "")).strip()
    print(
        f"[notify/upload] hash={body.file_hash!r}  "
        f"company={body.company_domain!r}  founder={email!r}"
    )
    return JSONResponse(content={"status": "ok"}, status_code=200)


_BUCKET_TO_VERTICAL: dict[str, str] = {
    "SAAS":  "SAAS",
    "LEND":  "FINTECH",
    "ECOM":  "MARKETPLACE",
    "INSUR": "INSURTECH",
    "OTH":   "GENERAL",
}


@app.get("/api/founder/config")
@limiter.limit("30/minute")
async def founder_config(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Auto-detects company_id and vertical for the authenticated founder.

    Derives company context from the JWT email domain so the Founder Portal
    never needs to ask the user to choose their company manually.

    Returns
    -------
    JSON ``{ "company_id", "vertical", "is_known", "domain" }``
    """
    email: str = (token.get("email") or token.get("sub", "")).strip()
    if not email or "@" not in email:
        raise HTTPException(status_code=422, detail="email no disponible en el token")
    domain = email.split("@", 1)[1].lower()
    comp_id, _, bucket_id, is_known = get_company_id(domain)
    vertical = _BUCKET_TO_VERTICAL.get(bucket_id, "GENERAL")
    return JSONResponse(content={
        "company_id": comp_id,
        "vertical":   vertical,
        "is_known":   is_known,
        "domain":     domain,
    })


class FinalizeRequest(BaseModel):
    """Body for POST /api/founder/finalize."""
    file_hashes:    list[str]
    company_domain: str
    file_names:     list[str] = []
    manual_kpis:    dict[str, str] | None = None


@app.post("/api/founder/finalize")
@limiter.limit("10/minute")
async def founder_finalize(
    request: Request,
    body: FinalizeRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Finalize a founder's expediente.

    Marks the submission set as complete and dispatches an HTML confirmation
    email to the founder.  Safe to call even when no email transport is
    configured — the dev fallback prints to stdout and the endpoint still
    returns 200.

    Parameters
    ----------
    body.file_hashes    : SHA-256 prefixes of every processed document.
    body.company_domain : Company slug, e.g. ``"solvento.com"``.
    body.file_names     : Display names of the processed files (for the email).
    body.manual_kpis    : Any KPI key/value pairs supplied manually by the founder.

    Returns
    -------
    JSON ``{ "status": "ok", "message": "...", "sent_to": email }``
    """
    from src.services.email_service import send_receipt_email
    from src.services.hash_service  import generate_vault_seal

    founder_email: str = (token.get("email") or token.get("sub", "")).strip()

    if not founder_email:
        raise HTTPException(status_code=403, detail="email no disponible en el token")

    # Derive company domain from token when not supplied
    company_domain = (body.company_domain or "").strip()
    if not company_domain:
        company_id_claim = (token.get("company_id") or "").strip()
        if "@" in company_id_claim:
            company_domain = company_id_claim.split("@")[-1]
        elif company_id_claim:
            company_domain = company_id_claim
        else:
            company_domain = founder_email.split("@")[-1] if "@" in founder_email else "cometa"

    file_names  = body.file_names or [f[:16] + "…" for f in body.file_hashes]
    manual_kpis = body.manual_kpis or {}
    processed_at = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M")
    period_id    = datetime.now(timezone.utc).strftime("%Y")

    # ── Vault Seal — SHA-256 integridad del expediente ────────────────────────
    # Cubre: company_id + file_hashes (ordenados) + timestamp de finalización.
    # Genera un fingerprint determinista y auditable que va al correo y a BQ.
    vault_seal = generate_vault_seal(
        company_id   = company_domain,
        file_hash    = body.file_hashes[0] if body.file_hashes else "",
        kpi_rows     = [
            {"kpi_key": k, "raw_value": v, "unit": "", "is_valid": True}
            for k, v in manual_kpis.items()
        ],
        processed_at = processed_at,
    )

    print(
        f"[founder/finalize] company={company_domain!r}  "
        f"files={len(body.file_hashes)}  founder={founder_email!r}  "
        f"seal={vault_seal[:16]}…"
    )

    # ── BigQuery — guardar recibo digital (non-fatal) ─────────────────────────
    try:
        insert_upload_log(
            company_id    = company_domain,
            founder_email = founder_email,
            vault_seal    = vault_seal,
            file_hashes   = body.file_hashes,
            manual_kpis   = manual_kpis if manual_kpis else None,
            period_id     = period_id,
        )
    except Exception as _log_err:
        print(f"⚠️  [founder/finalize] upload_log insert failed (non-fatal): {_log_err}")

    # ── Correo de confirmación con Sello de Bóveda ───────────────────────────
    send_receipt_email(
        to_email       = founder_email,
        company_domain = company_domain,
        period         = period_id,
        vault_seal     = vault_seal,
        file_hash      = body.file_hashes[0] if body.file_hashes else "",
        kpi_count      = len(manual_kpis),
        processed_at   = processed_at,
    )

    return JSONResponse(
        content={
            "status":     "ok",
            "message":    "Expediente registrado. Se ha enviado tu Recibo Digital al correo.",
            "sent_to":    founder_email,
            "vault_seal": vault_seal,
        },
        status_code=200,
    )


# ── Admin: invite founder ──────────────────────────────────────────────────────

_INVITE_EXPIRE_HOURS = 48  # noqa: F811 — already defined near setup-password; safe duplicate

_ADMIN_INVITE_FRONTEND_URL = os.getenv(
    "NEXTAUTH_URL",
    "https://cometa-vault-frontend-92572839783.us-central1.run.app",
)

_EMAIL_RE_ADMIN = re.compile(r"^[^\s@]+@[^\s@]+\.[^\s@]+$")


class AdminInviteRequest(BaseModel):
    """Body for POST /api/admin/invite — restricted to ANALISTA role."""
    email:        str
    company_name: str
    name:         str = ""


@app.get("/api/admin/invitations")
@limiter.limit("30/minute")
async def admin_invitations(
    request: Request,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Return all FOUNDER users with their activation status.
    Restricted to ANALISTA role.

    Returns
    -------
    JSON ``{ "invitations": [{ email, name, company_id, status }] }``
    """
    if token.get("role") not in ("ANALISTA",):
        raise HTTPException(status_code=403, detail="Solo analistas pueden ver invitaciones.")

    users = _load_users()
    founders = [
        {
            "email":      u.get("email", ""),
            "name":       u.get("name", ""),
            "company_id": u.get("company_id", ""),
            "status":     u.get("status", "ACTIVE"),
        }
        for u in users
        if u.get("role") == "FOUNDER"
    ]
    return JSONResponse(content={"invitations": founders})


@app.post("/api/admin/invite")
@limiter.limit("20/minute")
async def admin_invite(
    request: Request,
    body: AdminInviteRequest,
    token: dict = Depends(_require_auth),
) -> JSONResponse:
    """
    Create and dispatch a secure founder invitation (ANALISTA-only).

    Flow
    ----
    1. Enforce ANALISTA role.
    2. Validate email format.  Any existing record (ACTIVE or PENDING_INVITE)
       is dropped and recreated — this allows re-inviting founders who lost
       their setup link or whose access needs to be reset.
    3. Auto-derive role: @cometa.vc / @cometa.fund / @cometavc.com → ANALISTA,
       everything else → FOUNDER.
    4. Generate a signed JWT invite token (48 h, type="invite").
    5. Register PENDING_INVITE record in users.json atomically.
    6. Send invite email via email_service.send_invite_email().
    7. Return { status, email, company_name, setup_url }.
    """
    from src.services.email_service import send_invite_email  # lazy import

    # ── 1. Role guard ─────────────────────────────────────────────────────────
    caller_role = token.get("role", "")
    if caller_role not in ("ANALISTA",):
        raise HTTPException(status_code=403, detail="Solo analistas pueden enviar invitaciones.")

    email_lc = (body.email or "").strip().lower()
    if not _EMAIL_RE_ADMIN.match(email_lc):
        raise HTTPException(status_code=422, detail=f"Email inválido: {email_lc!r}")

    company_name = body.company_name.strip()
    if not company_name:
        raise HTTPException(status_code=422, detail="El nombre de la empresa es obligatorio.")

    # ── 2. Duplicate check — ACTIVE users are reset to PENDING_INVITE (re-invite) ─
    users = _load_users()
    existing = next((u for u in users if u.get("email", "").lower() == email_lc), None)
    # Both ACTIVE and PENDING_INVITE: drop stale record and recreate fresh
    if existing:
        users = [u for u in users if u.get("email", "").lower() != email_lc]

    # ── 3. Derive role from email domain ──────────────────────────────────────
    invite_role = "ANALISTA" if any(
        email_lc.endswith(d) for d in _INTERNAL_DOMAINS
    ) else "FOUNDER"

    # ── 4. Generate invite token ──────────────────────────────────────────────
    now = datetime.now(timezone.utc)
    invite_payload = {
        "type":         _INVITE_TOKEN_TYPE,
        "sub":          email_lc,
        "email":        email_lc,
        "company_name": company_name,
        "iat":          now,
        "exp":          now + timedelta(hours=_INVITE_EXPIRE_HOURS),
    }
    invite_token = jwt.encode(invite_payload, _JWT_SECRET, algorithm=_JWT_ALGORITHM)
    setup_url    = f"{_ADMIN_INVITE_FRONTEND_URL}/auth/setup-password?token={invite_token}"

    # ── 5. Register PENDING_INVITE user ──────────────────────────────────────
    company_domain = email_lc.split("@")[1] if "@" in email_lc else ""
    company_id     = company_domain or company_name.lower().replace(" ", "_")
    placeholder_pw = f"LOCKED:{secrets.token_hex(24)}"
    new_user_dict  = {
        "id":         generate_hybrid_id(email_lc),
        "email":      email_lc,
        "password":   placeholder_pw,
        "name":       body.name.strip() or company_name,
        "role":       invite_role,
        "company_id": company_id,
        "status":     "PENDING_INVITE",
    }

    all_users = users + [new_user_dict]
    validated: list[UserSchema] = [UserSchema.model_validate(u) for u in all_users]
    _save_users(validated)
    print(f"[admin/invite] Registered {email_lc!r} as PENDING_INVITE role={invite_role} (company={company_name!r})")

    # ── 5. Send invite email ──────────────────────────────────────────────────
    sent, email_error = send_invite_email(
        to_email=email_lc,
        company_name=company_name,
        setup_url=setup_url,
    )
    if sent:
        print(f"[admin/invite] Email sent to {email_lc!r}")
    else:
        print(f"[admin/invite] WARN: Email not sent — {email_error}")

    return JSONResponse(
        content={
            "status":       "ok",
            "email":        email_lc,
            "company_name": company_name,
            "setup_url":    setup_url,
            "email_sent":   sent,
            "email_error":  email_error,
        },
        status_code=200,
    )


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8080)))
